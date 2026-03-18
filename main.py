#!/usr/bin/env python3
"""
JIRA Analyzer - Simple GUI Version
Provides GUI for credential entry and dropdown selections
"""

import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, simpledialog
import threading
import time
import sys
import os
import json
import ssl
import urllib.request
import base64
import subprocess
import logging
from datetime import datetime

# Import the backend analyzer
sys.path.insert(0, os.path.dirname(__file__))
from jira_analyzer import JIRAAnalyzer, CredentialManager


class SimpleGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("BENTO - GUI")
        self.root.geometry("1400x750")
        
        # Initialize file logger
        self._init_file_logger()
        
        # Create analyzer with log callback to redirect output to GUI
        self.analyzer = JIRAAnalyzer(log_callback=self.log)
        self.repos = []
        self.branches = []
        self.config = self.load_config()
        
        # Chat continuation variables (initialized but not used until analysis starts)
        self.current_chat_messages = []
        self.chat_window = None
        
        # Workflow state management
        self.workflow_file = None
        self.workflow_state = {}
        
        # GUI lock state (must be initialized BEFORE create_widgets)
        self.gui_locked = False
        self.lockable_buttons = []
        
        self.create_widgets()
        
        self.check_saved_credentials()
        
        # Handle window closing
        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)
    
    def on_closing(self):
        """Handle window closing event"""
        if messagebox.askokcancel("Quit", "Do you want to close the application?"):
            # Close file logger
            if hasattr(self, 'file_logger') and self.file_logger:
                for handler in self.file_logger.handlers[:]:
                    handler.close()
                    self.file_logger.removeHandler(handler)
            self.root.destroy()
    
    def _init_file_logger(self):
        """Initialize file logger with timestamp in Logs folder"""
        logs_dir = "Logs"
        if not os.path.exists(logs_dir):
            os.makedirs(logs_dir)
        
        now = datetime.now()
        timestamp_file = now.strftime("%d%b%y")  # e.g., 03Mar26
        timestamp_time = now.strftime("%H%M")     # e.g., 1639
        log_filename = os.path.join(logs_dir, f"BentoLog_{timestamp_file}_{timestamp_time}hrs.log")
        
        # Create a dedicated logger for file output
        logger_id = now.strftime("%Y%m%d_%H%M%S")
        self.file_logger = logging.getLogger(f"bento_{logger_id}")
        self.file_logger.setLevel(logging.DEBUG)
        
        # Prevent propagation to root logger
        self.file_logger.propagate = False
        
        # File handler with detailed format
        file_handler = logging.FileHandler(log_filename, encoding='utf-8')
        file_handler.setLevel(logging.DEBUG)
        file_formatter = logging.Formatter(
            '%(asctime)s | %(levelname)-8s | %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
        file_handler.setFormatter(file_formatter)
        self.file_logger.addHandler(file_handler)
        
        # Store log file path for reference
        self.log_file_path = log_filename
        
        # Write header
        self.file_logger.info("=" * 80)
        self.file_logger.info("BENTO - Build, Evaluate, Navigate, Test & Orchestrate")
        self.file_logger.info(f"Session started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        self.file_logger.info(f"Log file: {log_filename}")
        self.file_logger.info("=" * 80)

    def load_config(self):
        """Load configuration from settings.json"""
        try:
            if os.path.exists('settings.json'):
                with open('settings.json', 'r') as f:
                    return json.load(f)
            return {}
        except Exception as e:
            print(f"Error loading config: {e}")
            return {}
    
    def load_modes_config(self):
        """Load AI modes configuration from ai_modes.json"""
        try:
            if os.path.exists('ai_modes.json'):
                with open('ai_modes.json', 'r') as f:
                    return json.load(f)
            return {}
        except Exception as e:
            print(f"Error loading AI modes: {e}")
            return {}
    
    def save_config(self):
        """Save configuration to settings.json and model configs to ai_modes.json"""
        try:
            # Load existing config to preserve other settings
            existing_config = self.load_config()
            
            config = {
                'jira': {
                    'base_url': self.jira_url_var.get(),
                    'project_key': self.jira_project_var.get()
                },
                'bitbucket': {
                    'base_url': self.bb_url_var.get(),
                    'project_key': self.project_key_var.get()
                },
                'model_gateway': {
                    'base_url': self.model_url_var.get()
                },
                'settings': existing_config.get('settings', {
                    'timeout': 300
                }),
                'compile': {
                    'last_tgz_label': self.tgz_label_var.get() if hasattr(self, 'tgz_label_var') else ""
                }
            }
            with open('settings.json', 'w') as f:
                json.dump(config, f, indent=2)
            
            # Save model configs to ai_modes.json (preserve existing modes and prompts)
            modes_data = self.load_modes_config()
            modes_data['models'] = {
                'analysis': {
                    'name': self.analysis_model_var.get(),
                    'temperature': modes_data.get('models', {}).get('analysis', {}).get('temperature', 0.7),
                    'max_tokens': modes_data.get('models', {}).get('analysis', {}).get('max_tokens', 2000)
                },
                'code_generation': {
                    'name': self.code_model_var.get(),
                    'temperature': modes_data.get('models', {}).get('code_generation', {}).get('temperature', 0.3),
                    'max_tokens': modes_data.get('models', {}).get('code_generation', {}).get('max_tokens', 4000)
                }
            }
            with open('ai_modes.json', 'w') as f:
                json.dump(modes_data, f, indent=2)
            
            self.log("✓ Configuration saved to settings.json")
            self.log("✓ Model configuration saved to ai_modes.json")
            self.log(f"  Analysis Model: {self.analysis_model_var.get()}")
            self.log(f"  Code Model: {self.code_model_var.get()}")
            return True
        except Exception as e:
            self.log(f"✗ Error saving config: {e}")
            return False
    
    def create_widgets(self):
        # Left side - Main container with tabs
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Title with DEBUG indicator
        title_frame = ttk.Frame(main_frame)
        title_frame.grid(row=0, column=0, columnspan=3, pady=10, sticky=(tk.W, tk.E))
        
        title = ttk.Label(title_frame, text="BENTO - Build, Evaluate, Navigate, Test & Orchestrate", font=('Arial', 16, 'bold'))
        title.pack(side=tk.LEFT)
        
        # DEBUG mode indicator (initially hidden)
        self.debug_indicator = ttk.Label(title_frame, text="🐛 DEBUG MODE",
                                         font=('Arial', 10, 'bold'),
                                         foreground='red',
                                         background='yellow',
                                         padding="5")
        # Don't pack it yet - will be shown when debug mode is enabled
        
        # Create Notebook (Tabs)
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.grid(row=1, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        # Right side - Log Section
        log_frame = ttk.LabelFrame(self.root, text="Progress Log", padding="5")
        log_frame.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 10), pady=10)
        
        self.log_text = scrolledtext.ScrolledText(log_frame, height=48, width=45)
        self.log_text.pack(fill=tk.BOTH, expand=True)
        
        # Create tabs - Reorganized to match new workflow
        # New workflow: 1. Fetch JIRA → 2. Analyze JIRA → 3. Clone repo → 4. Implementation → 5. Test Scenarios → 6. Validation + Risk
        self.create_home_tab()
        self.create_fetch_issue_tab()
        self.create_analyze_jira_tab()
        self.create_repo_tab()
        self.create_implementation_tab()
        self.create_test_tab()
        self.create_risk_tab()  # Will be renamed to Validation + Risk
        # Removed: create_impact_tab() - no longer in workflow
        
        # Configure grid weights
        self.root.columnconfigure(0, weight=0)  # Main content fixed at 600px
        self.root.columnconfigure(1, weight=1)  # Log takes remaining space
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(1, weight=1)
    
    def create_home_tab(self):
        """Create the home tab with full workflow automation"""
        home_tab = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(home_tab, text="🏠 Home")
        
        # Configuration Section (in home tab)
        config_frame = ttk.LabelFrame(home_tab, text="Configuration", padding="10")
        config_frame.grid(row=0, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        # Row 0: JIRA URL
        ttk.Label(config_frame, text="JIRA URL:").grid(row=0, column=0, sticky=tk.W, pady=2)
        self.jira_url_var = tk.StringVar(value=self.config.get('jira', {}).get('base_url', 'https://micron.atlassian.net'))
        ttk.Entry(config_frame, textvariable=self.jira_url_var, width=70).grid(row=0, column=1, columnspan=3, pady=2, sticky=tk.W)
        
        # Row 1: Bitbucket URL
        ttk.Label(config_frame, text="Bitbucket URL:").grid(row=1, column=0, sticky=tk.W, pady=2)
        self.bb_url_var = tk.StringVar(value=self.config.get('bitbucket', {}).get('base_url', 'https://bitbucket.micron.com/bbdc/scm'))
        ttk.Entry(config_frame, textvariable=self.bb_url_var, width=70).grid(row=1, column=1, columnspan=3, pady=2, sticky=tk.W)
        
        # Row 2: Project Keys
        ttk.Label(config_frame, text="Bitbucket Project:").grid(row=2, column=0, sticky=tk.W, pady=2)
        self.project_key_var = tk.StringVar(value=self.config.get('bitbucket', {}).get('project_key', 'TESTSSD'))
        ttk.Entry(config_frame, textvariable=self.project_key_var, width=27).grid(row=2, column=1, pady=2, sticky=tk.W)
        
        ttk.Label(config_frame, text="JIRA Project:").grid(row=2, column=2, sticky=tk.W, pady=2, padx=(20,0))
        self.jira_project_var = tk.StringVar(value=self.config.get('jira', {}).get('project_key', 'TSESSD'))
        ttk.Entry(config_frame, textvariable=self.jira_project_var, width=27).grid(row=2, column=3, pady=2, sticky=tk.W)
        
        # Row 3: Model Gateway URL
        ttk.Label(config_frame, text="Model Gateway:").grid(row=3, column=0, sticky=tk.W, pady=2)
        self.model_url_var = tk.StringVar(value=self.config.get('model_gateway', {}).get('base_url', 'https://model-gateway.gcldgenaigw.gc.micron.com/api/v1'))
        ttk.Entry(config_frame, textvariable=self.model_url_var, width=70).grid(row=3, column=1, columnspan=3, pady=2, sticky=tk.W)
        
        # Model variables (loaded from ai_modes.json, not displayed in GUI — edit ai_modes.json directly)
        modes_data = self.load_modes_config()
        models = modes_data.get('models', {})
        self.analysis_model_var = tk.StringVar(value=models.get('analysis', {}).get('name', 'gemini-3-pro-preview'))
        self.code_model_var = tk.StringVar(value=models.get('code_generation', {}).get('name', 'claude-sonnet-4-5'))
        
        # Row 4: Debug Checkbox
        self.debug_var = tk.BooleanVar(value=False)
        self.debug_var.trace_add('write', self.toggle_debug_mode)
        ttk.Checkbutton(config_frame, text="Enable Debug Mode", variable=self.debug_var).grid(row=4, column=0, columnspan=2, sticky=tk.W, pady=2)

        # Row 5: Save and Test buttons
        config_btn_frame = ttk.Frame(config_frame)
        config_btn_frame.grid(row=5, column=0, columnspan=4, pady=5)
        ttk.Button(config_btn_frame, text="Save Config", command=self.save_config).pack(side=tk.LEFT, padx=5)
        test_cfg_btn = ttk.Button(config_btn_frame, text="Test Config", command=self.test_config_with_credential_check)
        test_cfg_btn.pack(side=tk.LEFT, padx=5)
        self.lockable_buttons.append(test_cfg_btn)
        
        # Credentials Section (in home tab)
        cred_frame = ttk.LabelFrame(home_tab, text="Credentials (Encrypted)", padding="10")
        cred_frame.grid(row=1, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        ttk.Label(cred_frame, text="Email:").grid(row=0, column=0, sticky=tk.W, pady=2)
        self.email_var = tk.StringVar()
        ttk.Entry(cred_frame, textvariable=self.email_var, width=70).grid(row=0, column=1, columnspan=2, pady=2, sticky=tk.W)
        
        ttk.Label(cred_frame, text="JIRA Token:").grid(row=1, column=0, sticky=tk.W, pady=2)
        self.jira_token_var = tk.StringVar()
        self.jira_token_entry = ttk.Entry(cred_frame, textvariable=self.jira_token_var, show='*', width=70)
        self.jira_token_entry.grid(row=1, column=1, pady=2, sticky=tk.W)
        ttk.Button(cred_frame, text="👁", width=3,
                   command=lambda: self.toggle_password_visibility(self.jira_token_entry)).grid(row=1, column=2, padx=(2, 0), pady=2)
        
        ttk.Label(cred_frame, text="Bitbucket Token:").grid(row=2, column=0, sticky=tk.W, pady=2)
        self.bb_token_var = tk.StringVar()
        self.bb_token_entry = ttk.Entry(cred_frame, textvariable=self.bb_token_var, show='*', width=70)
        self.bb_token_entry.grid(row=2, column=1, pady=2, sticky=tk.W)
        ttk.Button(cred_frame, text="👁", width=3,
                   command=lambda: self.toggle_password_visibility(self.bb_token_entry)).grid(row=2, column=2, padx=(2, 0), pady=2)
        
        ttk.Label(cred_frame, text="Model API Key:").grid(row=3, column=0, sticky=tk.W, pady=2)
        self.model_key_var = tk.StringVar()
        self.model_key_entry = ttk.Entry(cred_frame, textvariable=self.model_key_var, show='*', width=70)
        self.model_key_entry.grid(row=3, column=1, pady=2, sticky=tk.W)
        ttk.Button(cred_frame, text="👁", width=3,
                   command=lambda: self.toggle_password_visibility(self.model_key_entry)).grid(row=3, column=2, padx=(2, 0), pady=2)
        
        btn_frame = ttk.Frame(cred_frame)
        btn_frame.grid(row=4, column=0, columnspan=3, pady=5)
        
        load_btn = ttk.Button(btn_frame, text="Load Credentials", command=self.load_credentials_with_lock)
        load_btn.pack(side=tk.LEFT, padx=2)
        self.lockable_buttons.append(load_btn)
        
        ttk.Button(btn_frame, text="Save", command=self.save_credentials).pack(side=tk.LEFT, padx=2)
        
        # Task Section (in home tab)
        task_frame = ttk.LabelFrame(home_tab, text="Task Details - Full Workflow", padding="10")
        task_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        # JIRA Issue with pre-filled prefix
        ttk.Label(task_frame, text="JIRA Issue:").grid(row=0, column=0, sticky=tk.W, pady=2)
        self.issue_var = tk.StringVar(value=f"{self.jira_project_var.get()}-")
        self.issue_var.trace_add('write', self.sync_issue_to_tabs)
        ttk.Entry(task_frame, textvariable=self.issue_var, width=70).grid(row=0, column=1, pady=2)
        
        ttk.Label(task_frame, text="Repository:").grid(row=1, column=0, sticky=tk.W, pady=2)
        self.repo_var = tk.StringVar()
        self.repo_var.trace_add('write', self.sync_repo_to_tabs)
        # Make combobox searchable by allowing typing
        self.repo_combo = ttk.Combobox(task_frame, textvariable=self.repo_var, width=67)
        self.repo_combo.grid(row=1, column=1, pady=2)
        self.repo_combo.bind('<<ComboboxSelected>>', self.on_repo_selected)
        self.repo_combo.bind('<KeyRelease>', self.filter_repos)
        
        ttk.Label(task_frame, text="Base Branch:").grid(row=2, column=0, sticky=tk.W, pady=2)
        self.branch_var = tk.StringVar()
        self.branch_var.trace_add('write', self.sync_branch_to_tabs)
        # Make combobox searchable by allowing typing
        self.branch_combo = ttk.Combobox(task_frame, textvariable=self.branch_var, width=67)
        self.branch_combo.grid(row=2, column=1, pady=2)
        self.branch_combo.bind('<KeyRelease>', self.filter_branches)
        
        # Feature Branch (auto-populated, editable)
        ttk.Label(task_frame, text="Feature Branch:").grid(row=3, column=0, sticky=tk.W, pady=2)
        self.feature_branch_var = tk.StringVar(value="If empty, will automatically be named as 'feature/TSESSD-XXXX'")
        self.feature_branch_entry = ttk.Entry(task_frame, textvariable=self.feature_branch_var, width=70)
        self.feature_branch_entry.grid(row=3, column=1, pady=2)
        
        # Workflow buttons (inside Task Details frame)
        workflow_btn_frame = ttk.Frame(task_frame)
        workflow_btn_frame.grid(row=4, column=0, columnspan=2, pady=10)
        ttk.Button(workflow_btn_frame, text="📂 Load Workflow", command=self.load_workflow_file).pack(side=tk.LEFT, padx=5)
        ttk.Button(workflow_btn_frame, text="🚀 Start Full Analysis Workflow", command=self.start_analysis,
                  style='Accent.TButton').pack(side=tk.LEFT, padx=5)

    
    
    def toggle_password_visibility(self, entry_widget):
        """Toggle between showing and hiding password in an entry widget"""
        if entry_widget.cget('show') == '*':
            entry_widget.configure(show='')
        else:
            entry_widget.configure(show='*')
    
    def init_workflow_file(self, issue_key):
        """Initialize workflow file for an issue"""
        # Check if we already have a loaded workflow file for this issue
        if self.workflow_file:
            basename = os.path.basename(self.workflow_file)
            if basename == f"{issue_key}_workflow.txt":
                # Already initialized for this issue, keep using existing file path
                # This supports legacy files in root directory
                self.load_workflow_state()
                return

        # Create Workflows directory if not exists
        workflows_dir = "Workflows"
        if not os.path.exists(workflows_dir):
            os.makedirs(workflows_dir)
            self.log(f"Created directory: {workflows_dir}")
            
        self.workflow_file = os.path.join(workflows_dir, f"{issue_key}_workflow.txt")
        self.load_workflow_state()
    
    def load_workflow_state(self):
        """Load workflow state from file"""
        if self.workflow_file and os.path.exists(self.workflow_file):
            try:
                with open(self.workflow_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                    
                # Parse the workflow file
                self.workflow_state = {}
                current_section = None
                current_content = []
                
                for line in content.split('\n'):
                    if line.startswith('=== ') and line.endswith(' ==='):
                        # Save previous section
                        if current_section:
                            self.workflow_state[current_section] = '\n'.join(current_content).strip()
                        # Start new section
                        current_section = line.replace('=== ', '').replace(' ===', '').strip()
                        current_content = []
                    elif current_section:
                        current_content.append(line)
                
                # Save last section
                if current_section:
                    self.workflow_state[current_section] = '\n'.join(current_content).strip()
                
                self.log(f"✓ Loaded workflow state from {self.workflow_file}")
                self.log(f"  Available sections: {', '.join(self.workflow_state.keys())}")
            except Exception as e:
                self.log(f"⚠ Error loading workflow state: {e}")
                self.workflow_state = {}
        else:
            self.workflow_state = {}
    
    def save_workflow_step(self, step_name, content, issue_key=None):
        """Save a workflow step to the consolidated file"""
        if issue_key:
            self.init_workflow_file(issue_key)
        
        if not self.workflow_file:
            self.log("⚠ No workflow file initialized")
            return
        
        # Update state
        self.workflow_state[step_name] = content
        
        # Write entire workflow file
        try:
            with open(self.workflow_file, 'w', encoding='utf-8') as f:
                f.write(f"JIRA WORKFLOW STATE FILE\n")
                f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"{'='*80}\n\n")
                
                for section, section_content in self.workflow_state.items():
                    f.write(f"=== {section} ===\n")
                    f.write(f"{section_content}\n\n")
            
            self.log(f"✓ Saved {step_name} to {self.workflow_file}")
        except Exception as e:
            self.log(f"✗ Error saving workflow step: {e}")
    
    def load_workflow_file(self):
        """Load a workflow file and populate fields"""
        from tkinter import filedialog
        
        # Determine initial directory
        initial_dir = os.path.join(os.getcwd(), "Workflows")
        if not os.path.exists(initial_dir):
            initial_dir = os.getcwd()
            
        # Open file dialog filtered for workflow files
        filename = filedialog.askopenfilename(
            title="Select Workflow File",
            filetypes=[("Workflow files", "*_workflow.txt"), ("All files", "*.*")],
            initialdir=initial_dir
        )
        
        if not filename:
            return
        
        try:
            # Set the workflow file
            self.workflow_file = filename
            self.load_workflow_state()
            
            # Extract issue key from filename
            issue_key = self.get_workflow_step("ISSUE_KEY")
            if not issue_key:
                # Try to extract from filename
                basename = os.path.basename(filename)
                if basename.endswith("_workflow.txt"):
                    issue_key = basename.replace("_workflow.txt", "")
            
            # Populate fields from workflow state
            if issue_key:
                self.issue_var.set(issue_key)
                self.log(f"✓ Loaded workflow for {issue_key}")
            
            # Get repository info
            repo_info = self.get_workflow_step("REPOSITORY_INFO")
            repo_path = None
            if repo_info:
                # Parse repository info
                for line in repo_info.split('\n'):
                    if line.startswith("Repository:"):
                        repo_name = line.split(":", 1)[1].strip()
                        self.repo_var.set(repo_name)
                    elif line.startswith("Base branch:"):
                        base_branch = line.split(":", 1)[1].strip()
                        self.branch_var.set(base_branch)
                    elif line.startswith("Feature branch:"):
                        feature_branch = line.split(":", 1)[1].strip()
                        self.feature_branch_var.set(feature_branch)
                    elif line.startswith("Local path:"):
                        repo_path = line.split(":", 1)[1].strip()
            
            # Auto-populate repository path to Implementation and Test Scenarios tabs
            if repo_path:
                self.impl_repo_var.set(repo_path)
                self.log(f"  Auto-populated repo path: {repo_path}")
            
            # Show summary
            sections = list(self.workflow_state.keys())
            messagebox.showinfo(
                "Workflow Loaded",
                f"Loaded workflow for {issue_key}\n\n"
                f"Available sections:\n" + "\n".join(f"  • {s}" for s in sections)
            )
            
        except Exception as e:
            self.log(f"✗ Error loading workflow file: {e}")
            messagebox.showerror("Error", f"Failed to load workflow file:\n{e}")
    
    
    def get_workflow_step(self, step_name):
        """Get a workflow step from state"""
        return self.workflow_state.get(step_name)
    
    def ensure_prerequisite(self, step_name, prerequisite_step, auto_run_func):
        """Ensure a prerequisite step is completed, auto-run if missing"""
        if not self.get_workflow_step(prerequisite_step):
            self.log(f"⚠ {prerequisite_step} is required for {step_name}")
            self.log(f"  Automatically running {prerequisite_step}...")
            
            response = messagebox.askyesno(
                "Prerequisite Required",
                f"{prerequisite_step} is required for {step_name}.\n\n"
                f"Would you like to run it automatically?"
            )
            
            if response:
                return auto_run_func()
            else:
                messagebox.showwarning("Cannot Proceed", f"Cannot proceed without {prerequisite_step}")
                return None
        
        return self.get_workflow_step(prerequisite_step)
    
    def create_fetch_issue_tab(self):
        """Tab for fetching JIRA issue only"""
        tab = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(tab, text="📋 Fetch Issue")
        
        ttk.Label(tab, text="Fetch JIRA Issue", font=('Arial', 14, 'bold')).grid(row=0, column=0, columnspan=2, pady=10)
        
        # Issue input
        ttk.Label(tab, text="JIRA Issue Key:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.fetch_issue_var = tk.StringVar(value=f"{self.jira_project_var.get()}-")
        ttk.Entry(tab, textvariable=self.fetch_issue_var, width=40).grid(row=1, column=1, pady=5)
        
        # Fetch button
        ttk.Button(tab, text="Fetch Issue", command=self.fetch_issue_only_with_lock).grid(row=2, column=0, columnspan=2, pady=10)
        
        # Result display
        result_frame = ttk.LabelFrame(tab, text="Issue Details", padding="10")
        result_frame.grid(row=3, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        self.issue_result_text = scrolledtext.ScrolledText(result_frame, height=20, width=70, wrap=tk.WORD)
        self.issue_result_text.pack(fill=tk.BOTH, expand=True)
        
        tab.columnconfigure(1, weight=1)
        tab.rowconfigure(3, weight=1)
    
    def create_analyze_jira_tab(self):
        """Tab for analyzing JIRA with AI"""
        tab = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(tab, text="🤖 Analyze JIRA")
        
        ttk.Label(tab, text="Analyze JIRA Request with AI", font=('Arial', 14, 'bold')).grid(row=0, column=0, columnspan=2, pady=10)
        
        # Issue input
        ttk.Label(tab, text="JIRA Issue Key:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.analyze_issue_var = tk.StringVar(value=f"{self.jira_project_var.get()}-")
        ttk.Entry(tab, textvariable=self.analyze_issue_var, width=40).grid(row=1, column=1, pady=5)
        
        # Analyze button
        ttk.Button(tab, text="Analyze with AI", command=self.analyze_jira_only_with_lock).grid(row=2, column=0, columnspan=2, pady=10)
        
        # Result display
        result_frame = ttk.LabelFrame(tab, text="AI Analysis Result", padding="10")
        result_frame.grid(row=3, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        self.analyze_result_text = scrolledtext.ScrolledText(result_frame, height=20, width=70, wrap=tk.WORD)
        self.analyze_result_text.pack(fill=tk.BOTH, expand=True)
        
        tab.columnconfigure(1, weight=1)
        tab.rowconfigure(3, weight=1)
    
    def create_repo_tab(self):
        """Tab for repository operations"""
        tab = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(tab, text="📦 Repository")
        
        ttk.Label(tab, text="Clone Repository & Create Branch", font=('Arial', 14, 'bold')).grid(row=0, column=0, columnspan=2, pady=10)
        
        # Repository selection
        ttk.Label(tab, text="Repository:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.repo_tab_var = tk.StringVar()
        repo_combo = ttk.Combobox(tab, textvariable=self.repo_tab_var, width=37)
        repo_combo.grid(row=1, column=1, pady=5)
        
        # Branch selection
        ttk.Label(tab, text="Base Branch:").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.branch_tab_var = tk.StringVar()
        ttk.Entry(tab, textvariable=self.branch_tab_var, width=40).grid(row=2, column=1, pady=5)
        
        # Issue key for feature branch
        ttk.Label(tab, text="Issue Key:").grid(row=3, column=0, sticky=tk.W, pady=5)
        self.repo_issue_var = tk.StringVar(value=f"{self.jira_project_var.get()}-")
        ttk.Entry(tab, textvariable=self.repo_issue_var, width=40).grid(row=3, column=1, pady=5)
        
        # Buttons Frame
        btn_frame = ttk.Frame(tab)
        btn_frame.grid(row=4, column=0, columnspan=2, pady=10)
        
        # Clone button
        ttk.Button(btn_frame, text="1. Clone Repository", command=self.clone_repo_action_with_lock).pack(side=tk.LEFT, padx=5)
        
        # Create Feature Branch button
        ttk.Button(btn_frame, text="2. Create Feature Branch", command=self.create_feature_branch_action_with_lock).pack(side=tk.LEFT, padx=5)
        
        # Result display
        result_frame = ttk.LabelFrame(tab, text="Repository Status", padding="10")
        result_frame.grid(row=5, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        self.repo_result_text = scrolledtext.ScrolledText(result_frame, height=15, width=70, wrap=tk.WORD)
        self.repo_result_text.pack(fill=tk.BOTH, expand=True)
        
        tab.columnconfigure(1, weight=1)
        tab.rowconfigure(5, weight=1)
    
    def create_impact_tab(self):
        """Tab for code impact analysis"""
        tab = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(tab, text="🔍 Impact Analysis")
        
        ttk.Label(tab, text="Analyze Code Impact", font=('Arial', 14, 'bold')).grid(row=0, column=0, columnspan=2, pady=10)
        
        # Repository path
        ttk.Label(tab, text="Repository Path:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.impact_repo_var = tk.StringVar()
        ttk.Entry(tab, textvariable=self.impact_repo_var, width=40).grid(row=1, column=1, pady=5)
        
        # Analyze button
        ttk.Button(tab, text="Analyze Impact", command=self.analyze_impact_only_with_lock).grid(row=2, column=0, columnspan=2, pady=10)
        
        # Result display
        result_frame = ttk.LabelFrame(tab, text="Impact Analysis Result", padding="10")
        result_frame.grid(row=3, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        self.impact_result_text = scrolledtext.ScrolledText(result_frame, height=20, width=70, wrap=tk.WORD)
        self.impact_result_text.pack(fill=tk.BOTH, expand=True)
        
        tab.columnconfigure(1, weight=1)
        tab.rowconfigure(3, weight=1)
    
    def create_test_tab(self):
        """Tab for test scenario generation"""
        tab = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(tab, text="🧪 Test Scenarios")
        
        ttk.Label(tab, text="Generate Test Scenarios", font=('Arial', 14, 'bold')).grid(row=0, column=0, columnspan=3, pady=10)
        
        # Issue Key (auto-populated from home tab)
        ttk.Label(tab, text="JIRA Issue Key:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.test_issue_var = tk.StringVar(value=f"{self.jira_project_var.get()}-")
        ttk.Entry(tab, textvariable=self.test_issue_var, width=50).grid(row=1, column=1, pady=5, sticky=(tk.W, tk.E))
        ttk.Label(tab, text="(Auto-populated from Home tab)", font=('Arial', 8), foreground='gray').grid(row=1, column=2, sticky=tk.W, padx=5)
        
        # Info label
        ttk.Label(tab, text="Test scenarios will be generated from workflow data",
                 font=('Arial', 9), foreground='gray').grid(row=2, column=0, columnspan=3, pady=5)
        
        # Generate button
        ttk.Button(tab, text="Generate Test Scenarios", command=self.generate_tests_only_with_lock).grid(row=3, column=0, columnspan=3, pady=10)
        
        # Result display
        result_frame = ttk.LabelFrame(tab, text="Test Scenarios", padding="10")
        result_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        self.test_result_text = scrolledtext.ScrolledText(result_frame, height=20, width=70, wrap=tk.WORD)
        self.test_result_text.pack(fill=tk.BOTH, expand=True)
        
        tab.columnconfigure(1, weight=1)
        tab.rowconfigure(4, weight=1)
    
    def create_risk_tab(self):
        """Tab for validation document and risk assessment (combined)"""
        tab = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(tab, text="📋 Validation & Risk")
        
        ttk.Label(tab, text="Generate Validation Document & Risk Assessment", font=('Arial', 14, 'bold')).grid(row=0, column=0, columnspan=3, pady=10)
        
        # Issue Key (auto-populated from home tab)
        ttk.Label(tab, text="JIRA Issue Key:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.risk_issue_var = tk.StringVar(value=f"{self.jira_project_var.get()}-")
        ttk.Entry(tab, textvariable=self.risk_issue_var, width=50).grid(row=1, column=1, pady=5, sticky=(tk.W, tk.E))
        ttk.Label(tab, text="(Auto-populated from Home tab)", font=('Arial', 8), foreground='gray').grid(row=1, column=2, sticky=tk.W, padx=5)
        
        # Info label
        ttk.Label(tab, text="Risk assessment will be generated from workflow data",
                 font=('Arial', 9), foreground='gray').grid(row=2, column=0, columnspan=3, pady=5)
        
        # Generate button
        ttk.Button(tab, text="Generate Validation & Risk Assessment", command=self.assess_risks_only_with_lock).grid(row=3, column=0, columnspan=3, pady=10)
        
        # Result display
        result_frame = ttk.LabelFrame(tab, text="Validation Document & Risk Assessment", padding="10")
        result_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        self.risk_result_text = scrolledtext.ScrolledText(result_frame, height=22, width=70, wrap=tk.WORD)
        self.risk_result_text.pack(fill=tk.BOTH, expand=True)
        
        tab.columnconfigure(1, weight=1)
        tab.rowconfigure(4, weight=1)
    
    # ================================================================
    # WATCHER HEALTH MONITOR
    # ================================================================
    def build_watcher_health_panel(self, parent_frame):
        """
        Watcher Health Monitor panel.
        Shows live status of RAW_ZIP, RELEASE_TGZ, watcher lock, last ZIP status.
        Auto-refreshes every 30 seconds.
        """
        panel = ttk.LabelFrame(parent_frame, text="🔍 Watcher Health Monitor", padding=8)
        panel.pack(fill=tk.X, padx=10, pady=5)

        rows = {}

        def _add_row(label_text, default="Checking..."):
            row = ttk.Frame(panel)
            row.pack(fill=tk.X, pady=2)
            ttk.Label(row, text=label_text, width=28, anchor="w",
                      font=("Arial", 9, "bold")).pack(side=tk.LEFT)
            val_lbl = ttk.Label(row, text=default, anchor="w", foreground="gray")
            val_lbl.pack(side=tk.LEFT)
            return val_lbl

        rows["raw_zip"]     = _add_row("📂 RAW_ZIP Folder:")
        rows["release_tgz"] = _add_row("📂 RELEASE_TGZ Folder:")
        rows["watcher"]     = _add_row("🤖 Watcher Process:")

        # Container for recent builds
        recent_frame = ttk.Frame(panel)
        recent_frame.pack(fill=tk.X, pady=(5, 2))
        ttk.Label(recent_frame, text="📊 Recent Builds:", width=28, anchor="w", font=("Segoe UI", 9, "bold")).pack(side=tk.LEFT, anchor=tk.N)
        
        # Text widget for multi-line colored status
        bg_color = ttk.Style().lookup('TFrame', 'background') or '#f0f0f0'
        builds_text = tk.Text(recent_frame, height=6, width=65, bg=bg_color, relief="flat", font=("Segoe UI", 9))
        builds_text.pack(side=tk.LEFT, fill=tk.X, expand=True)
        builds_text.tag_config("success", foreground="#28a745")
        builds_text.tag_config("failed", foreground="#dc3545")
        builds_text.tag_config("in_progress", foreground="#fd7e14")
        builds_text.tag_config("timeout", foreground="purple")
        builds_text.tag_config("unknown", foreground="gray")
        builds_text.tag_config("default", foreground="black")

        ttk.Button(
            panel, text="🔄 Refresh Now",
            command=lambda: _refresh()
        ).pack(anchor="e", padx=5, pady=3)

        def _refresh():
            try:
                if not panel.winfo_exists():
                    return
            except Exception:
                return
            raw_zip_folder     = self.raw_zip_var.get().strip()     if hasattr(self, "raw_zip_var")     else r"P:\temp\BENTO\RAW_ZIP"
            release_tgz_folder = self.release_tgz_var.get().strip() if hasattr(self, "release_tgz_var") else r"P:\temp\BENTO\RELEASE_TGZ"
            repo_dir           = r"C:\BENTO\adv_ibir_master"

            # RAW_ZIP reachability
            if os.path.isdir(raw_zip_folder):
                rows["raw_zip"].config(text="✅ Reachable  " + raw_zip_folder, foreground="#28a745")
            else:
                rows["raw_zip"].config(text="❌ NOT REACHABLE: " + raw_zip_folder, foreground="#dc3545")

            # RELEASE_TGZ reachability
            if os.path.isdir(release_tgz_folder):
                rows["release_tgz"].config(text="✅ Reachable  " + release_tgz_folder, foreground="#28a745")
            else:
                rows["release_tgz"].config(text="❌ NOT REACHABLE: " + release_tgz_folder, foreground="#dc3545")

            # Watcher lock freshness
            lock_path = os.path.join(repo_dir, ".bento_build_lock")
            if os.path.exists(lock_path):
                age = int(time.time() - os.path.getmtime(lock_path))
                rows["watcher"].config(text="🔒 Build Lock active (age=" + str(age) + "s)", foreground="#fd7e14")
            else:
                has_lock = False
                if os.path.isdir(raw_zip_folder):
                    for f in os.listdir(raw_zip_folder):
                        if f.endswith(".bento_lock"):
                            has_lock = True
                            break
                if has_lock:
                    rows["watcher"].config(text="🟡 ZIP lock present (picking up?)", foreground="#fd7e14")
                else:
                    rows["watcher"].config(text="✅ Idle (no active build lock)", foreground="#28a745")

            # Last ZIP and status — scan .bento_status files directly
            # (successful ZIPs are deleted after compile, status files persist)
            try:
                all_status_files = []
                if os.path.isdir(raw_zip_folder):
                    for f in os.listdir(raw_zip_folder):
                        if f.endswith(".bento_status"):
                            all_status_files.append(os.path.join(raw_zip_folder, f))

                if all_status_files:
                    newest = max(all_status_files, key=os.path.getmtime)
                    zip_name = os.path.basename(newest).replace(".bento_status", "")
                    rows["last_zip"].config(text=zip_name, foreground="black")
                    with open(newest) as sf:
                        sdata  = json.load(sf)
                        state  = sdata.get("status", "unknown")
                        detail = sdata.get("detail", "")
                    colour_map = {"success": "green", "failed": "red",
                                  "in_progress": "orange", "timeout": "purple"}
                    rows["last_status"].config(
                        text=state.upper() + " — " + detail[:60],
                        foreground=colour_map.get(state, "gray")
                    )
                else:
                    # No status files — check for ZIPs not yet picked up
                    zips = [f for f in os.listdir(raw_zip_folder)
                            if f.endswith(".zip") and ".bento_" not in f
                            ] if os.path.isdir(raw_zip_folder) else []
                    if zips:
                        newest_zip = max(zips, key=lambda f: os.path.getmtime(
                            os.path.join(raw_zip_folder, f)))
                        rows["last_zip"].config(text=newest_zip, foreground="gray")
                        rows["last_status"].config(
                            text="Waiting for watcher...", foreground="orange")
                    else:
                        rows["last_zip"].config(text="(none)", foreground="gray")
                        rows["last_status"].config(text="—", foreground="gray")
            except Exception as e:
                builds_text.config(state="normal")
                builds_text.delete("1.0", tk.END)
                builds_text.insert(tk.END, f"Error: {str(e)}", "failed")
                builds_text.config(state="disabled")

            panel.after(30000, _refresh)

        _refresh()
        return panel

    # ================================================================
    # REAL-TIME BUILD STATUS MONITOR
    # ================================================================
    def open_build_status_monitor(self, issue_key, hostname, env, raw_zip_folder):
        """
        Opens a live build status monitor window.
        Polls the .bento_status sidecar on the shared RAW_ZIP folder every 3s.
        Uses local closure flag (not instance var) so multiple monitors are safe.
        """
        mon_win = tk.Toplevel(self.root)
        mon_win.title("Build Monitor — " + hostname + " (" + env + ")")
        self._centre_dialog(mon_win, 620, 400)
        mon_win.resizable(False, False)

        ttk.Label(mon_win, text="Live Build Status",
                  font=("Arial", 12, "bold")).pack(anchor=tk.W, padx=16, pady=(14, 2))
        ttk.Label(mon_win,
                  text=hostname + "  (" + env + ")  —  JIRA: " + issue_key,
                  font=("Arial", 9), foreground="#555555").pack(anchor=tk.W, padx=16)
        ttk.Separator(mon_win, orient="horizontal").pack(fill=tk.X, padx=16, pady=8)

        badge_var = tk.StringVar(value="🟡  Waiting for watcher to pick up ZIP...")
        badge_lbl = ttk.Label(mon_win, textvariable=badge_var,
                               font=("Arial", 10, "bold"), foreground="orange")
        badge_lbl.pack(anchor=tk.W, padx=16, pady=(0, 6))

        detail_var = tk.StringVar(value="")
        ttk.Label(mon_win, textvariable=detail_var,
                  font=("Courier", 9), foreground="#333333",
                  wraplength=580, justify=tk.LEFT).pack(anchor=tk.W, padx=16)

        elapsed_var = tk.StringVar(value="")
        ttk.Label(mon_win, textvariable=elapsed_var,
                  font=("Arial", 8), foreground="gray").pack(anchor=tk.W, padx=16, pady=(4, 0))

        ttk.Separator(mon_win, orient="horizontal").pack(fill=tk.X, padx=16, pady=8)
        ttk.Label(mon_win,
                  text="Build log on tester:  C:\\BENTO\\logs\\build_<ZIP_NAME>.log",
                  font=("Arial", 8), foreground="#777777").pack(anchor=tk.W, padx=16)
        ttk.Button(mon_win, text="Close", command=mon_win.destroy).pack(pady=12)

        # Fix 2: local mutable flag — safe if multiple monitors open simultaneously
        active = [True]
        start_time = time.time()

        colour_map = {"success": "green", "failed": "red",
                      "in_progress": "orange", "pending": "orange", "timeout": "purple"}
        icon_map   = {"success": "✅", "failed": "❌",
                      "in_progress": "🔨", "pending": "🟡", "timeout": "⏱"}

        def _poll():
            # Fix 4: stop if window destroyed
            try:
                if not mon_win.winfo_exists():
                    return
            except Exception:
                return
            if not active[0]:
                return

            elapsed_var.set("Elapsed: " + str(int(time.time() - start_time)) + "s")

            try:
                if os.path.isdir(raw_zip_folder):
                    for fname in os.listdir(raw_zip_folder):
                        if not fname.endswith(".bento_status"):
                            continue
                        base = fname.replace(".zip.bento_status", "")
                        if (hostname.upper() in base.upper() and
                                env.upper() in base.upper() and
                                issue_key.upper() in base.upper()):
                            with open(os.path.join(raw_zip_folder, fname)) as f:
                                data = json.load(f)
                            state  = data.get("status", "unknown")
                            detail = data.get("detail", "")
                            badge_var.set(icon_map.get(state, "🔵") + "  " +
                                          state.upper().replace("_", " "))
                            badge_lbl.config(foreground=colour_map.get(state, "gray"))
                            detail_var.set(detail[:200] if detail else "")
                            if state in ("success", "failed", "timeout"):
                                return
                            break
                    else:
                        badge_var.set("🟡  Waiting for watcher to pick up ZIP...")
                        badge_lbl.config(foreground="orange")
            except Exception:
                pass

            mon_win.after(3000, _poll)

        def _on_close():
            active[0] = False
            mon_win.destroy()

        mon_win.protocol("WM_DELETE_WINDOW", _on_close)
        _poll()

    # ================================================================
    # PER-TESTER LIVE STATUS BADGES
    # ================================================================
    def open_multi_compile_status(self, targets):
        """
        Opens a live status window showing per-tester build progress.
        targets: list of (hostname, env) tuples.
        Returns (callbacks dict, status_win).
        """
        status_win = tk.Toplevel(self.root)
        status_win.title("Multi-Tester Compile Status")
        w, h = 600, 80 + len(targets) * 55
        status_win.resizable(False, False)
        self._centre_dialog(status_win, w, h)

        ttk.Label(status_win, text="Live Compile Status",
                  font=("Arial", 12, "bold")).pack(pady=10)

        rows = {}

        for hostname, env in targets:
            key = (hostname, env)
            row_frame = ttk.Frame(status_win, relief="groove", padding=6)
            row_frame.pack(fill=tk.X, padx=15, pady=4)

            ttk.Label(row_frame, text=hostname + "  (" + env + ")",
                      width=25, anchor="w",
                      font=("Courier", 10, "bold")).pack(side=tk.LEFT)

            status_lbl = ttk.Label(row_frame, text="🟡 Waiting...",
                                   foreground="orange", width=22, anchor="w")
            status_lbl.pack(side=tk.LEFT, padx=10)

            time_lbl = ttk.Label(row_frame, text="0s",
                                 foreground="gray", width=8, anchor="e")
            time_lbl.pack(side=tk.RIGHT)

            rows[key] = {"label": status_lbl, "time_label": time_lbl,
                         "start_time": None, "running": False}

        def _tick():
            try:
                if not status_win.winfo_exists():
                    return
            except Exception:
                return
            for key, row in rows.items():
                if row["start_time"] and row.get("running", False):
                    elapsed = int(time.time() - row["start_time"])
                    row["time_label"].config(text=str(elapsed) + "s")
            status_win.after(1000, _tick)

        _tick()

        def make_update(key):
            def update_status(state, elapsed=None):
                row = rows[key]
                lbl = row["label"]
                if state == "building":
                    row["start_time"] = time.time()
                    row["running"]    = True
                    lbl.config(text="🟡 Building...", foreground="orange")
                elif state == "success":
                    row["running"] = False
                    t = str(elapsed) + "s" if elapsed else ""
                    lbl.config(text="✅ Done " + t, foreground="green")
                    row["time_label"].config(foreground="green")
                elif state == "failed":
                    row["running"] = False
                    lbl.config(text="❌ Failed", foreground="red")
                    row["time_label"].config(foreground="red")
                elif state == "timeout":
                    row["running"] = False
                    lbl.config(text="⏱ Timeout", foreground="purple")
                status_win.update_idletasks()
            return update_status

        callbacks = {key: make_update(key) for key in rows}
        return callbacks, status_win

    def create_implementation_tab(self):
        """Tab for code implementation"""
        tab = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(tab, text="💻 Implementation")
        
        # ── SUB-NOTEBOOK ──
        impl_notebook = ttk.Notebook(tab)
        impl_notebook.pack(fill=tk.BOTH, expand=True)

        # ==========================================
        # SUB-TAB 1: AI Plan Generator
        # ==========================================
        ai_tab = ttk.Frame(impl_notebook, padding="10")
        impl_notebook.add(ai_tab, text="🧠 AI Plan Generator")
        
        ttk.Label(ai_tab, text="Generate Implementation Plan", font=('Arial', 14, 'bold')).grid(row=0, column=0, columnspan=3, pady=10)
        
        # JIRA Issue Key (auto-populated from home tab)
        ttk.Label(ai_tab, text="JIRA Issue Key:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.impl_issue_var = tk.StringVar(value=f"{self.jira_project_var.get()}-")
        ttk.Entry(ai_tab, textvariable=self.impl_issue_var, width=50).grid(row=1, column=1, pady=5, sticky=(tk.W, tk.E))
        ttk.Label(ai_tab, text="(Auto-populated from Home tab)", font=('Arial', 8), foreground='gray').grid(row=1, column=2, sticky=tk.W, padx=5)
        
        # Repository path (local cloned repo)
        ttk.Label(ai_tab, text="Local Repo Path:").grid(row=2, column=0, sticky=tk.W, pady=5)
        repo_path_frame = ttk.Frame(ai_tab)
        repo_path_frame.grid(row=2, column=1, pady=5, sticky=(tk.W, tk.E))
        self.impl_repo_var = tk.StringVar()
        ttk.Entry(repo_path_frame, textvariable=self.impl_repo_var, width=47).pack(side=tk.LEFT)
        ttk.Button(repo_path_frame, text="📁", width=3,
                   command=lambda: self._browse_directory(self.impl_repo_var, "Select Repository Folder")).pack(side=tk.LEFT, padx=(2, 0))
        ttk.Label(ai_tab, text="(Path to cloned repository for AI indexing)", font=('Arial', 8), foreground='gray').grid(row=2, column=2, sticky=tk.W, padx=5)
        
        # Generate button
        ttk.Button(ai_tab, text="Generate Implementation Plan", command=self.generate_implementation_only_with_lock).grid(row=3, column=0, columnspan=3, pady=10)

        # ── Implementation Plan Result ──────────────────────────────
        result_frame = ttk.LabelFrame(ai_tab, text="Implementation Plan", padding="10")
        result_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)

        self.impl_result_text = scrolledtext.ScrolledText(result_frame, height=25, width=70, wrap=tk.WORD)
        self.impl_result_text.pack(fill=tk.BOTH, expand=True)

        ai_tab.columnconfigure(1, weight=1)
        ai_tab.rowconfigure(4, weight=1)

        # ==========================================
        # SUB-TAB 2: TP Compilation & Health
        # ==========================================
        compile_subtab = ttk.Frame(impl_notebook, padding="10")
        impl_notebook.add(compile_subtab, text="📦 TP Compilation & Health")

        # ── Compile TP Package Section ──────────────────────────────
        compile_frame = ttk.LabelFrame(compile_subtab, text="Compile TP Package", padding="10")
        compile_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=5)

        # ── TESTER REGISTRY ──
        # Persisted to bento_testers.json next to main.py.
        # In-memory dict: "HOSTNAME (ENV)" -> (hostname, env)
        self._TESTER_REGISTRY = {}
        self._tester_registry_file = os.path.join(
            os.path.dirname(os.path.abspath(__file__)), "bento_testers.json"
        )
        self._load_tester_registry()   # populate from JSON or defaults

        compile_frame.columnconfigure(0, weight=1)
        compile_frame.columnconfigure(1, weight=1)

        # ── 1. TARGET TESTERS FRAME ──
        targets_frame = ttk.LabelFrame(compile_frame, text="1. Target Testers", padding="10")
        targets_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5), pady=(0, 5))
        targets_frame.columnconfigure(0, weight=1)

        # Search bar
        search_frame = ttk.Frame(targets_frame)
        search_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 5))
        
        ttk.Label(search_frame, text="Search:").pack(side=tk.LEFT, padx=(0, 5))
        self.tester_search_var = tk.StringVar()
        self._tester_search_entry = ttk.Entry(search_frame, textvariable=self.tester_search_var, width=15)
        self._tester_search_entry.pack(side=tk.LEFT)
        self._tester_search_entry.insert(0, "Search...")
        self._tester_search_entry.config(foreground="gray")
        
        ttk.Label(search_frame, text="(Shift+Click for multi)", font=("Arial", 8, "italic"), foreground="gray").pack(side=tk.LEFT, padx=(5, 0))

        # Listbox and buttons
        list_frame = ttk.Frame(targets_frame)
        list_frame.grid(row=1, column=0, sticky=(tk.W, tk.E))
        
        self._tester_listbox = tk.Listbox(
            list_frame, 
            selectmode=tk.MULTIPLE,
            height=4,
            width=28,
            exportselection=False
        )
        self._tester_listbox.pack(side=tk.LEFT, fill=tk.Y)
        
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self._tester_listbox.yview)
        scrollbar.pack(side=tk.LEFT, fill=tk.Y)
        self._tester_listbox.config(yscrollcommand=scrollbar.set)
        
        btn_frame_list = ttk.Frame(list_frame)
        btn_frame_list.pack(side=tk.LEFT, padx=(5, 0), anchor=tk.N)
        
        ttk.Button(btn_frame_list, text="+ Add Tester", width=12, command=self._open_add_tester_dialog).pack(pady=(0, 5))
        ttk.Button(btn_frame_list, text="🗑 Remove", width=12, command=self._remove_selected_tester).pack(pady=(0, 5))

        # Selection info
        info_frame = ttk.Frame(targets_frame)
        info_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=(5, 0))
        
        ttk.Label(info_frame, text="Selected:", font=("Arial", 9, "bold")).pack(side=tk.LEFT, padx=(0, 5))
        self._compile_mode_var = tk.StringVar(value="No tester selected")
        self._compile_mode_lbl = ttk.Label(info_frame, textvariable=self._compile_mode_var, font=("Arial", 9, "italic"), foreground="#cc0000", wraplength=250)
        self._compile_mode_lbl.pack(side=tk.LEFT)

        # ── 2. CONFIGURATION & PATHS FRAME ──
        config_frame = ttk.LabelFrame(compile_frame, text="2. Configuration & Paths", padding="10")
        config_frame.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(5, 0), pady=(0, 5))
        config_frame.columnconfigure(1, weight=1)
        
        # TGZ Label
        ttk.Label(config_frame, text="TGZ Label:").grid(row=0, column=0, sticky=tk.W, pady=3, padx=(0, 5))
        last_label = self.config.get('compile', {}).get('last_tgz_label', '')
        self.tgz_label_var = tk.StringVar(value=last_label)
        ttk.Entry(config_frame, textvariable=self.tgz_label_var, width=22).grid(row=0, column=1, sticky=tk.W, pady=3)
        ttk.Label(config_frame, text="(blank = default)", font=("Arial", 8), foreground="gray").grid(row=0, column=2, sticky=tk.W, padx=2)

        # RAW_ZIP Path
        ttk.Label(config_frame, text="RAW_ZIP:").grid(row=1, column=0, sticky=tk.W, pady=3, padx=(0, 5))
        raw_zip_frame = ttk.Frame(config_frame)
        raw_zip_frame.grid(row=1, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=3)
        self.raw_zip_var = tk.StringVar(value=r"P:\temp\BENTO\RAW_ZIP")
        ttk.Entry(raw_zip_frame, textvariable=self.raw_zip_var, width=28).pack(side=tk.LEFT)
        ttk.Button(raw_zip_frame, text="📁", width=3, command=lambda: self._browse_directory(self.raw_zip_var, "Select RAW_ZIP Folder")).pack(side=tk.LEFT, padx=(2, 0))

        # RELEASE_TGZ Path
        ttk.Label(config_frame, text="RELEASE_TGZ:").grid(row=2, column=0, sticky=tk.W, pady=3, padx=(0, 5))
        release_tgz_frame = ttk.Frame(config_frame)
        release_tgz_frame.grid(row=2, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=3)
        self.release_tgz_var = tk.StringVar(value=r"P:\temp\BENTO\RELEASE_TGZ")
        ttk.Entry(release_tgz_frame, textvariable=self.release_tgz_var, width=28).pack(side=tk.LEFT)
        ttk.Button(release_tgz_frame, text="📁", width=3, command=lambda: self._browse_directory(self.release_tgz_var, "Select RELEASE_TGZ Folder")).pack(side=tk.LEFT, padx=(2, 0))

        # ── 3. ACTION FRAME ──
        action_frame = ttk.Frame(compile_frame)
        action_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(5, 5))
        
        btn_container = ttk.Frame(action_frame)
        btn_container.pack(expand=True)
        
        # Style for prominent compile button
        style = ttk.Style()
        style.configure('Compile.TButton', font=('Arial', 10, 'bold'), padding=6)
        
        self.compile_btn = ttk.Button(
            btn_container, text="🚀 Compile on Selected Tester(s)",
            style='Compile.TButton',
            command=self.trigger_compile_with_lock,
            width=35
        )
        self.compile_btn.pack(pady=(0, 2))
        
        self.compile_status_var = tk.StringVar(value="")
        self.compile_status_lbl = ttk.Label(btn_container, textvariable=self.compile_status_var, font=("Arial", 9, "bold"), foreground="#0066cc")
        self.compile_status_lbl.pack()

        # ── Wire up search + listbox ──
        def _search_changed(*_):
            q = self.tester_search_var.get().strip().lower()
            if q and q != "search...":
                matches = [k for k in self._TESTER_REGISTRY if q in k.lower()]
            else:
                matches = list(self._TESTER_REGISTRY.keys())
            
            # Remember current selections
            current_selections = [self._tester_listbox.get(i) for i in self._tester_listbox.curselection()]
            
            # Update listbox
            self._tester_listbox.delete(0, tk.END)
            for item in matches:
                self._tester_listbox.insert(tk.END, item)
            
            # Restore selections if they're still in the filtered list
            for i, item in enumerate(matches):
                if item in current_selections:
                    self._tester_listbox.selection_set(i)
            
            self._refresh_tester_mode()

        def _search_focus_in(event):
            if self._tester_search_entry.get() == "Search...":
                self._tester_search_entry.delete(0, tk.END)
                self._tester_search_entry.config(foreground="black")

        def _search_focus_out(event):
            if not self._tester_search_entry.get():
                self._tester_search_entry.insert(0, "Search...")
                self._tester_search_entry.config(foreground="gray")

        self.tester_search_var.trace_add("write", _search_changed)
        self._tester_search_entry.bind("<FocusIn>",  _search_focus_in)
        self._tester_search_entry.bind("<FocusOut>", _search_focus_out)
        self._tester_listbox.bind("<<ListboxSelect>>",
                                lambda e: self._refresh_tester_mode())

        # Populate dropdown and set initial state
        self._refresh_tester_dropdown()

        # ── Watcher Health Monitor — wrapped in grid row 1 ──────────
        health_wrapper = ttk.Frame(compile_subtab)
        health_wrapper.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(5, 0))
        health_wrapper.columnconfigure(0, weight=1)
        self.build_watcher_health_panel(health_wrapper)

        # ── Compile History — wrapped in grid row 2 ──────────────────
        history_wrapper = ttk.Frame(compile_subtab)
        history_wrapper.grid(row=2, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        history_wrapper.columnconfigure(0, weight=1)
        self._build_compile_history_section(history_wrapper)

        compile_subtab.columnconfigure(0, weight=1)
        compile_subtab.rowconfigure(2, weight=1)
    
    # Individual step methods with GUI locking wrappers
    def fetch_issue_only_with_lock(self):
        """Wrapper for fetch_issue_only with GUI locking"""
        self.lock_gui()
        try:
            return self.fetch_issue_only()
        finally:
            self.unlock_gui()
    
    def fetch_issue_only(self):
        """Fetch JIRA issue only"""
        issue_key = self.fetch_issue_var.get().strip().upper()
        if not issue_key:
            messagebox.showerror("Error", "Please enter JIRA issue key")
            return None
        
        # Initialize workflow file
        self.init_workflow_file(issue_key)
        
        self.log(f"\n[Fetching JIRA Issue: {issue_key}]")
        issue_data = self.analyzer.fetch_jira_issue(issue_key)
        
        if issue_data:
            fields = issue_data.get('fields', {})
            summary = fields.get('summary', 'N/A')
            
            # Get extracted fields (now done in fetch_jira_issue)
            extracted = issue_data.get('extracted_fields', {})
            
            # Build structured display with all extracted fields
            result = f"{'='*50}\n"
            result += f"JIRA ISSUE: {issue_key}\n"
            result += f"{'='*50}\n\n"
            
            result += f"SUMMARY:\n{summary}\n\n"
            
            result += f"{'─'*50}\n"
            result += f"WORK TYPE:\n{extracted.get('work_type', 'N/A')}\n\n"
            
            result += f"{'─'*50}\n"
            result += f"REPORTER:\n{extracted.get('reporter', 'N/A')}\n\n"
            
            result += f"{'─'*50}\n"
            result += f"ASSIGNEE:\n{extracted.get('assignee', 'Unassigned')}\n\n"
            
            result += f"{'─'*50}\n"
            result += f"DESCRIPTION:\n{extracted.get('description', 'N/A')}\n\n"
            
            result += f"{'─'*50}\n"
            result += f"COMPONENTS:\n{extracted.get('components_str', 'None')}\n\n"
            
            result += f"{'─'*50}\n"
            result += f"ACCEPTANCE CRITERIA:\n{extracted.get('acceptance_criteria', 'N/A')}\n\n"
            
            result += f"{'─'*50}\n"
            result += f"CHANGE CATEGORY:\n{extracted.get('change_category', 'N/A')}\n\n"
            
            result += f"{'─'*50}\n"
            result += f"ISSUE LINKS:\n{extracted.get('issue_links_count', 0)} linked issue(s)\n\n"
            
            result += f"{'─'*50}\n"
            result += f"ATTACHMENTS:\n{extracted.get('attachments_count', 0)} file(s)\n\n"
            
            result += f"{'─'*50}\n"
            result += f"RECENT COMMENTS (with roles):\n{extracted.get('comments_text', 'None')}\n"
            
            result += f"{'='*50}\n"
            
            # Save to workflow file
            self.save_workflow_step("ISSUE_KEY", issue_key)
            self.save_workflow_step("JIRA_ISSUE_DATA", result)
            
            self.issue_result_text.delete('1.0', tk.END)
            self.issue_result_text.insert('1.0', result)
            self.log(f"✓ Issue fetched and fields extracted successfully")
            return issue_data
        else:
            self.log(f"✗ Failed to fetch issue")
            messagebox.showerror("Error", f"Failed to fetch issue {issue_key}")
            return None
    
    def analyze_jira_only_with_lock(self):
        """Wrapper for analyze_jira_only with GUI locking"""
        self.lock_gui()
        try:
            return self.analyze_jira_only()
        finally:
            self.unlock_gui()
    
    def analyze_jira_only(self):
        """Analyze JIRA with AI only"""
        issue_key = self.analyze_issue_var.get().strip().upper()
        if not issue_key:
            messagebox.showerror("Error", "Please enter JIRA issue key")
            return None
        
        # Initialize workflow file
        self.init_workflow_file(issue_key)
        
        # Check if we already have JIRA analysis in workflow - if so, skip re-analysis
        existing_analysis = self.get_workflow_step("JIRA_ANALYSIS")
        if existing_analysis:
            self.log(f"✓ Found existing JIRA analysis in workflow file")
            self.log(f"  Opening interactive chat with existing analysis...")
            
            # Open chat with existing analysis
            self.root.after(100, lambda: self.create_interactive_chat(
                issue_key,
                "JIRA Analysis",
                existing_analysis,
                lambda: self.finalize_jira_analysis(issue_key, existing_analysis)
            ))
            
            return {'success': True, 'analysis': existing_analysis}
        
        # Check if we have issue data in workflow, if not fetch it
        issue_data_text = self.get_workflow_step("JIRA_ISSUE_DATA")
        if not issue_data_text:
            self.log("⚠ JIRA issue data not found in workflow, fetching from JIRA...")
            self.fetch_issue_var.set(issue_key)
            issue_data = self.fetch_issue_only()
            if not issue_data:
                return None
        else:
            self.log("✓ Found JIRA issue data in workflow file")
            # Re-fetch to get full data object (needed for analysis)
            issue_data = self.analyzer.fetch_jira_issue(issue_key)
        
        self.log(f"\n[Analyzing JIRA Issue with AI: {issue_key}]")
        
        if issue_data:
            jira_analysis = self.analyzer.analyze_jira_request(issue_data)
            if jira_analysis.get('success'):
                analysis_text = jira_analysis.get('analysis', '')
                
                # Open interactive chat FIRST for user to refine the analysis
                # Only save and mark complete when user approves
                self.log(f"✓ Initial analysis complete - opening interactive chat...")
                self.root.after(100, lambda: self.create_interactive_chat(
                    issue_key,
                    "JIRA Analysis",
                    analysis_text,
                    lambda: self.finalize_jira_analysis(issue_key, analysis_text)
                ))
                
                return jira_analysis
            else:
                self.log(f"✗ Analysis failed: {jira_analysis.get('error')}")
                messagebox.showerror("Error", f"Analysis failed: {jira_analysis.get('error')}")
                return None
        else:
            self.log(f"✗ Failed to fetch issue")
            messagebox.showerror("Error", f"Failed to fetch issue {issue_key}")
            return None
    
    def finalize_jira_analysis(self, issue_key, analysis_text):
        """Finalize JIRA analysis after user approval from interactive chat"""
        # Save to workflow file
        self.save_workflow_step("JIRA_ANALYSIS", analysis_text)
        
        # Update result window
        self.analyze_result_text.delete('1.0', tk.END)
        self.analyze_result_text.insert('1.0', analysis_text)
        
        # Log completion
        self.log(f"✓ JIRA analysis approved and saved")
        self.log(f"✓ Analysis for {issue_key} is complete")
    
    def clone_repo_action_with_lock(self):
        """Wrapper for clone_repo_action with GUI locking"""
        self.lock_gui()
        try:
            return self.clone_repo_action()
        finally:
            self.unlock_gui()
    
    def clone_repo_action(self):
        """Clone repository only"""
        repo_name = self.repo_tab_var.get().strip()
        branch_name = self.branch_tab_var.get().strip()
        issue_key = self.repo_issue_var.get().strip().upper()
        
        if not all([repo_name, branch_name, issue_key]):
            messagebox.showerror("Error", "All fields required")
            return None
        
        # Initialize workflow file
        self.init_workflow_file(issue_key)
        
        # Extract repo slug if needed
        if ' - ' in repo_name:
            repo_slug = repo_name.split(' - ')[0].strip()
        else:
            repo_slug = repo_name
        
        # Check if local repository already exists
        repos_dir = "Repos"
        expected_repo_path = os.path.join(repos_dir, f"{issue_key}_{repo_slug}")
        
        if os.path.exists(expected_repo_path):
            # Check if it's a git repository
            if os.path.exists(os.path.join(expected_repo_path, '.git')):
                warning_msg = f"⚠️ WARNING: Local repository already exists!\n\n"
                warning_msg += f"Path: {expected_repo_path}\n"
                warning_msg += "This may indicate:\n"
                warning_msg += "• You've already cloned this repository for this issue\n"
                warning_msg += "• There may be uncommitted changes\n"
                warning_msg += "• Previous work may be overwritten\n\n"
                warning_msg += "Please check the existing repository before proceeding.\n"
                warning_msg += "Consider using 'Load Workflow' if you want to continue previous work."
                
                self.log(f"⚠️ Local repository already exists: {expected_repo_path}")
                
                messagebox.showwarning("Repository Already Exists", warning_msg)
                return None
            else:
                # Directory exists but not a git repo
                warning_msg = f"⚠️ WARNING: Directory already exists but is not a git repository!\n\n"
                warning_msg += f"Path: {expected_repo_path}\n\n"
                warning_msg += "Please remove or rename this directory before cloning."
                
                self.log(f"⚠️ Directory exists but not a git repo: {expected_repo_path}")
                messagebox.showwarning("Directory Already Exists", warning_msg)
                return None
        
        self.log("\n=========start clone===========")
        self.log(f"[Cloning Repository: {repo_slug}]")
        repo_path = self.analyzer.clone_repository(repo_slug, branch_name, issue_key)
        self.log("============end clone==========")
        
        if repo_path:
            result = f"Repository cloned successfully!\n\n"
            result += f"Local path: {repo_path}\n"
            result += f"Repository: {repo_slug}\n"
            result += f"Base branch: {branch_name}\n"
            
            # Save to workflow file
            self.save_workflow_step("REPOSITORY_PATH", repo_path)
            self.save_workflow_step("REPOSITORY_INFO", result)
            
            self.repo_result_text.delete('1.0', tk.END)
            self.repo_result_text.insert('1.0', result)
            self.log(f"✓ Repository cloned and saved to workflow")
            return repo_path
        else:
            self.log(f"✗ Failed to clone repository")
            messagebox.showerror("Error", "Failed to clone repository")
            return None

    def create_feature_branch_action_with_lock(self):
        """Wrapper for create_feature_branch_action with GUI locking"""
        self.lock_gui()
        try:
            return self.create_feature_branch_action()
        finally:
            self.unlock_gui()

    def create_feature_branch_action(self):
        """Create feature branch only"""
        repo_name = self.repo_tab_var.get().strip()
        branch_name = self.branch_tab_var.get().strip()
        issue_key = self.repo_issue_var.get().strip().upper()
        
        # Check if we have a custom feature branch from the home tab
        feature_branch_input = self.feature_branch_var.get().strip()
        
        if not all([repo_name, branch_name, issue_key]):
            messagebox.showerror("Error", "All fields required")
            return None
            
        # Extract repo slug if needed
        if ' - ' in repo_name:
            repo_slug = repo_name.split(' - ')[0].strip()
        else:
            repo_slug = repo_name
            
        # Get repo path from workflow or construct it
        repo_path = self.get_workflow_step("REPOSITORY_PATH")
        if not repo_path:
            repo_path = os.path.join("Repos", f"{issue_key}_{repo_slug}")
            
        if not os.path.exists(repo_path):
             messagebox.showerror("Error", f"Repository path not found: {repo_path}\nPlease clone the repository first.")
             return None

        # Determine feature branch name
        if feature_branch_input and "if empty," not in feature_branch_input.lower() and feature_branch_input.lower() != "auto populate":
            feature_branch_name = feature_branch_input
        else:
            feature_branch_name = f"feature/{issue_key}"
            
        if self.analyzer.create_feature_branch(repo_path, issue_key, branch_name, feature_branch_input):
            result = f"Feature branch created successfully!\n\n"
            result += f"Local path: {repo_path}\n"
            result += f"Feature branch: {feature_branch_name}\n"
            result += f"Repository: {repo_slug}\n"
            result += f"Base branch: {branch_name}\n"
            
            # Update feature branch field in GUI
            self.feature_branch_var.set(feature_branch_name)
            
            # Save to workflow file
            self.save_workflow_step("REPOSITORY_PATH", repo_path)
            self.save_workflow_step("REPOSITORY_INFO", result)
            
            self.repo_result_text.delete('1.0', tk.END)
            self.repo_result_text.insert('1.0', result)
            self.log(f"✓ Feature branch created and saved to workflow")
            self.log(f"  Feature branch: {feature_branch_name}")
            return True
        else:
            self.log(f"✗ Failed to create feature branch")
            messagebox.showerror("Error", "Failed to create feature branch")
            return False
    
    def analyze_impact_only_with_lock(self):
        """Wrapper for analyze_impact_only with GUI locking"""
        self.lock_gui()
        try:
            return self.analyze_impact_only()
        finally:
            self.unlock_gui()
    
    def analyze_impact_only(self):
        """Analyze code impact only"""
        repo_path = self.impact_repo_var.get().strip()
        if not repo_path:
            messagebox.showerror("Error", "Please enter repository path")
            return
        
        self.log(f"\n[Analyzing Code Impact]")
        repo_index = self.analyzer.index_repository(repo_path)
        
        # Need JIRA analysis for context - prompt user
        issue_key = simpledialog.askstring("JIRA Issue", "Enter JIRA issue key for context:")
        if not issue_key:
            return
        
        issue_data = self.analyzer.fetch_jira_issue(issue_key)
        if issue_data:
            jira_analysis = self.analyzer.analyze_jira_request(issue_data)
            if jira_analysis.get('success'):
                impact_analysis = self.analyzer.analyze_code_impact(repo_path, repo_index, jira_analysis)
                if impact_analysis.get('success'):
                    self.impact_result_text.delete('1.0', tk.END)
                    self.impact_result_text.insert('1.0', impact_analysis.get('impact_analysis', ''))
                    self.log(f"✓ Impact analysis completed")
                else:
                    self.log(f"✗ Impact analysis failed")
                    messagebox.showerror("Error", "Impact analysis failed")
    
    def generate_tests_only_with_lock(self):
        """Wrapper for generate_tests_only with GUI locking"""
        self.lock_gui()
        try:
            return self.generate_tests_only()
        finally:
            self.unlock_gui()
    def generate_tests_only(self):
        """Generate test scenarios only"""
        issue_key = self.test_issue_var.get().strip().upper()
        
        if not issue_key or issue_key == f"{self.jira_project_var.get()}-":
            messagebox.showerror("Error", "Please enter JIRA issue key")
            return
        
        self.log(f"\n[Generating Test Scenarios for {issue_key}]")
        
        # Initialize workflow file to check for existing data
        self.init_workflow_file(issue_key)
        
        # Get repository path from workflow
        repo_path = self.get_workflow_step("REPOSITORY_PATH")
        if not repo_path:
            messagebox.showerror("Error",
                f"Repository path not found in workflow.\n\n"
                f"Please run the full analysis workflow first or clone the repository.")
            return
        
        # Check if we already have test scenarios in workflow
        existing_test_scenarios = self.get_workflow_step("TEST_SCENARIOS")
        if existing_test_scenarios:
            self.log(f"✓ Found existing test scenarios in workflow file")
            self.test_result_text.delete('1.0', tk.END)
            self.test_result_text.insert('1.0', existing_test_scenarios)
            self.log(f"✓ Test scenarios loaded from workflow")
            messagebox.showinfo("Success", "Test scenarios loaded from workflow file!")
            return
        
        # Index repository
        repo_index = self.analyzer.index_repository(repo_path)
        
        # Check for existing JIRA analysis in workflow
        jira_analysis_text = self.get_workflow_step("JIRA_ANALYSIS")
        if jira_analysis_text:
            self.log("✓ Using existing JIRA analysis from workflow file")
            jira_analysis = {'success': True, 'analysis': jira_analysis_text}
        else:
            self.log("⚠ JIRA analysis not found in workflow, fetching and analyzing...")
            issue_data = self.analyzer.fetch_jira_issue(issue_key)
            if not issue_data:
                messagebox.showerror("Error", f"Failed to fetch JIRA issue {issue_key}")
                return
            
            jira_analysis = self.analyzer.analyze_jira_request(issue_data)
            if not jira_analysis.get('success'):
                messagebox.showerror("Error", f"Failed to analyze JIRA issue: {jira_analysis.get('error')}")
                return
            
            # Save to workflow for future use
            self.save_workflow_step("JIRA_ANALYSIS", jira_analysis.get('analysis', ''))
        
        # Generate test scenarios
        test_scenarios = self.analyzer.generate_test_scenarios(jira_analysis, repo_index)
        if test_scenarios.get('success'):
            test_scenarios_text = test_scenarios.get('test_scenarios', '')
            
            # Open interactive chat for test scenarios review
            self.log(f"✓ Test scenarios generated - opening interactive chat...")
            self.root.after(100, lambda: self.create_interactive_chat(
                issue_key,
                "Test Scenarios",
                test_scenarios_text,
                lambda: self.finalize_test_scenarios(issue_key, test_scenarios_text)
            ))
        else:
            self.log(f"✗ Test generation failed: {test_scenarios.get('error')}")
            messagebox.showerror("Error", f"Test generation failed: {test_scenarios.get('error')}")
    
    def finalize_test_scenarios(self, issue_key, test_scenarios_text):
        """Finalize test scenarios after user approval from interactive chat"""
        # Save to workflow
        self.save_workflow_step("TEST_SCENARIOS", test_scenarios_text)
        
        # Display results
        self.test_result_text.delete('1.0', tk.END)
        self.test_result_text.insert('1.0', test_scenarios_text)
        
        # Log completion
        self.log(f"✓ Test scenarios approved and saved to workflow")
        self.log(f"✓ Test scenarios for {issue_key} are complete")
    
    def assess_risks_only_with_lock(self):
        """Wrapper for assess_risks_only with GUI locking"""
        self.lock_gui()
        try:
            return self.assess_risks_only()
        finally:
            self.unlock_gui()
    def assess_risks_only(self):
        """Generate validation document from template using workflow data"""
        # Get issue key from the risk tab field
        issue_key = self.risk_issue_var.get().strip().upper()
        
        if not issue_key or issue_key == f"{self.jira_project_var.get()}-":
            messagebox.showerror("Error", "Please enter JIRA issue key")
            return
        
        self.log(f"\n[Generating Validation Document for {issue_key}]")
        
        # Initialize workflow file to check for existing data
        self.init_workflow_file(issue_key)
        
        # Check for template file
        template_file = "template_validation.docx"
        if not os.path.exists(template_file):
            messagebox.showerror("Error",
                f"Template file not found: {template_file}\n\n"
                f"Please ensure the template file exists in the current directory.")
            return
        
        self.log(f"✓ Found template: {template_file}")
        
        # Get repository path from workflow
        repo_path = self.get_workflow_step("REPOSITORY_PATH")
        if not repo_path:
            messagebox.showerror("Error",
                f"Repository path not found in workflow.\n\n"
                f"Please run the full analysis workflow first or clone the repository.")
            return
        
        # Populate template with workflow data
        self.populate_validation_template(issue_key, template_file, repo_path)
        
        # Index repository
        repo_index = self.analyzer.index_repository(repo_path)
        
        # Check for existing JIRA analysis in workflow
        jira_analysis_text = self.get_workflow_step("JIRA_ANALYSIS")
        if jira_analysis_text:
            self.log("✓ Using existing JIRA analysis from workflow file")
            jira_analysis = {'success': True, 'analysis': jira_analysis_text}
        else:
            self.log("⚠ JIRA analysis not found in workflow, fetching and analyzing...")
            issue_data = self.analyzer.fetch_jira_issue(issue_key)
            if not issue_data:
                messagebox.showerror("Error", f"Failed to fetch JIRA issue {issue_key}")
                return
            
            jira_analysis = self.analyzer.analyze_jira_request(issue_data)
            if not jira_analysis.get('success'):
                messagebox.showerror("Error", f"Failed to analyze JIRA issue: {jira_analysis.get('error')}")
                return
            
            # Save to workflow for future use
            self.save_workflow_step("JIRA_ANALYSIS", jira_analysis.get('analysis', ''))
        
        # Check for existing impact analysis in workflow
        impact_analysis_text = self.get_workflow_step("IMPACT_ANALYSIS")
        if impact_analysis_text:
            self.log("✓ Using existing impact analysis from workflow file")
            impact_analysis = {'success': True, 'impact_analysis': impact_analysis_text}
        else:
            self.log("⚠ Impact analysis not found in workflow, generating...")
            impact_analysis = self.analyzer.analyze_code_impact(repo_path, repo_index, jira_analysis)
            if not impact_analysis.get('success'):
                messagebox.showerror("Error", f"Failed to analyze code impact: {impact_analysis.get('error')}")
                return
            
            # Save to workflow for future use
            self.save_workflow_step("IMPACT_ANALYSIS", impact_analysis.get('impact_analysis', ''))
        
        # Generate risk assessment
        risk_assessment = self.analyzer.assess_risks(jira_analysis, impact_analysis)
        if risk_assessment.get('success'):
            risk_assessment_text = risk_assessment.get('risk_assessment', '')
            
            # Open interactive chat for risk assessment review
            self.log(f"✓ Risk assessment generated - opening interactive chat...")
            self.root.after(100, lambda: self.create_interactive_chat(
                issue_key,
                "Validation & Risk Assessment",
                risk_assessment_text,
                lambda: self.finalize_risk_assessment(issue_key, risk_assessment_text)
            ))
        else:
            self.log(f"✗ Risk assessment failed: {risk_assessment.get('error')}")
            messagebox.showerror("Error", f"Risk assessment failed: {risk_assessment.get('error')}")
    
    def populate_validation_template(self, issue_key, template_file, repo_path):
        """Populate validation template with workflow data"""
        try:
            from docx import Document
        except ImportError:
            messagebox.showerror("Error",
                "python-docx library not installed.\n\n"
                "Please install it using: pip install python-docx")
            return
        
        try:
            self.log("  Reading template file...")
            doc = Document(template_file)
            
            # Extract template structure for AI context
            template_sections = []
            for para in doc.paragraphs:
                if para.text.strip():
                    template_sections.append(para.text)
            
            template_structure = '\n'.join(template_sections)
            self.log(f"  Template has {len(doc.paragraphs)} paragraphs")
            
            # Get workflow data
            jira_analysis = self.get_workflow_step("JIRA_ANALYSIS") or "Not available"
            impact_analysis = self.get_workflow_step("IMPACT_ANALYSIS") or "Not available"
            test_scenarios = self.get_workflow_step("TEST_SCENARIOS") or "Not available"
            risk_assessment = self.get_workflow_step("RISK_ASSESSMENT") or "Not available"
            repo_info = self.get_workflow_step("REPOSITORY_INFO") or "Not available"
            
            # Index repository for additional context
            self.log("  Indexing repository...")
            repo_index = self.analyzer.index_repository(repo_path)
            
            # Get validation_document prompt template from config
            prompt_template = self.analyzer.ai_client.get_prompt_template('validation_document')
            
            # Prepare repository statistics
            repo_stats = f"- Total Files: {repo_index['stats']['total_files']}\n"
            repo_stats += f"- File Types: {', '.join([f'{ext}({count})' for ext, count in sorted(repo_index['stats']['by_extension'].items(), key=lambda x: x[1], reverse=True)[:5]])}"
            
            # Format the prompt with actual values
            if prompt_template:
                populate_prompt = prompt_template.format(
                    template_structure=template_structure,
                    jira_analysis=jira_analysis[:2000],
                    impact_analysis=impact_analysis[:2000],
                    test_scenarios=test_scenarios[:2000],
                    risk_assessment=risk_assessment[:2000],
                    repo_info=repo_info,
                    repo_stats=repo_stats
                )
            else:
                # Fallback if prompt not in config
                populate_prompt = f"""You are populating a validation document template with information from a JIRA workflow.

**Template Structure:**
{template_structure}

**Available Workflow Data:**

**JIRA Analysis:**
{jira_analysis[:1000]}...

**Impact Analysis:**
{impact_analysis[:1000]}...

**Test Scenarios:**
{test_scenarios[:1000]}...

**Risk Assessment:**
{risk_assessment[:1000]}...

**Repository Info:**
{repo_info}

**Repository Statistics:**
{repo_stats}

Populate the template, leaving validation result sections for user to fill after checkout.
"""
            
            self.log("  Sending to AI for content generation...")
            messages = [{"role": "user", "content": populate_prompt}]
            
            # Use code_generation model for document generation
            result = self.analyzer.ai_client.chat_completion(messages, task_type="code_generation", mode="validation_document")
            
            if result['success']:
                populated_content = result['response']['choices'][0]['message']['content']
                self.log("✓ AI generated document content")
                
                # Clear existing paragraphs and add new content
                for para in doc.paragraphs:
                    para.clear()
                
                # Add populated content
                for line in populated_content.split('\n'):
                    doc.add_paragraph(line)
                
                # Save with issue key in filename
                output_file = f"{issue_key}_validation.docx"
                doc.save(output_file)
                
                self.log(f"✓ Validation document saved: {output_file}")
                
                # Display summary in result window
                summary = f"Validation document generated successfully!\n\n"
                summary += f"Output file: {output_file}\n\n"
                summary += f"Document populated with:\n"
                summary += f"- JIRA Analysis\n"
                summary += f"- Impact Analysis\n"
                summary += f"- Test Scenarios\n"
                summary += f"- Risk Assessment\n"
                summary += f"- Repository Information\n\n"
                summary += f"Preview:\n{populated_content[:500]}..."
                
                self.risk_result_text.delete('1.0', tk.END)
                self.risk_result_text.insert('1.0', summary)
                
                messagebox.showinfo("Success",
                    f"Validation document generated!\n\n"
                    f"Saved as: {output_file}")
                
            else:
                self.log(f"✗ AI generation failed: {result['error']}")
                messagebox.showerror("Error", f"Failed to generate content:\n{result['error']}")
                
        except Exception as e:
            self.log(f"✗ Error populating template: {str(e)}")
            import traceback
            traceback.print_exc()
            messagebox.showerror("Error", f"Failed to populate template:\n{str(e)}")
    
    def finalize_risk_assessment(self, issue_key, risk_assessment_text):
        """Finalize risk assessment after user approval from interactive chat"""
        # Save to workflow
        self.save_workflow_step("RISK_ASSESSMENT", risk_assessment_text)
        
        # Display results
        self.risk_result_text.delete('1.0', tk.END)
        self.risk_result_text.insert('1.0', risk_assessment_text)
        
        # Log completion
        self.log(f"✓ Risk assessment approved and saved to workflow")
        self.log(f"✓ Risk assessment for {issue_key} is complete")
    

    # ================================================================
    # COMPILE TP PACKAGE — TESTER REGISTRY + SELECTION
    # ================================================================

    _DEFAULT_TESTERS = {
        "IBIR-0383 (ABIT)":    ("IBIR-0383",    "ABIT", r"C:\BENTO\adv_ibir_master", "make release"),
        "MPT3HVM-0156 (SFN2)": ("MPT3HVM-0156", "SFN2", r"C:\BENTO\adv_ibir_master", "make release"),
        "CTOWTST-0031 (CNFG)": ("CTOWTST-0031", "CNFG", r"C:\BENTO\adv_ibir_master", "make release_supermicro"),
    }

    def _load_tester_registry(self):
        """Load tester list from bento_testers.json, falling back to defaults."""
        import json
        self._TESTER_REGISTRY = dict(self._DEFAULT_TESTERS)
        try:
            if os.path.exists(self._tester_registry_file):
                with open(self._tester_registry_file, "r") as f:
                    saved = json.load(f)
                for key, val in saved.items():
                    # Support both old format (2-tuple) and new format (4-tuple)
                    if isinstance(val, list):
                        if len(val) == 4:
                            self._TESTER_REGISTRY[key] = tuple(val)
                        elif len(val) == 2:
                            # Migrate old format: add default repo_dir and build_cmd
                            self._TESTER_REGISTRY[key] = tuple(val) + (r"C:\BENTO\adv_ibir_master", "make release")
        except Exception as e:
            self.log("[WARN] Could not load tester registry: " + str(e))

    def _save_tester_registry(self):
        """Persist tester list to bento_testers.json (both local and shared)."""
        import json
        try:
            # Convert to new format for local file
            data = {k: list(v) for k, v in self._TESTER_REGISTRY.items()}
            with open(self._tester_registry_file, "w") as f:
                json.dump(data, f, indent=2)
            
            # Also write to shared folder in watcher-compatible format
            shared_path = r"P:\temp\BENTO\bento_testers.json"
            try:
                # Ensure directory exists
                os.makedirs(os.path.dirname(shared_path), exist_ok=True)
                
                shared_data = {}
                for key, val in self._TESTER_REGISTRY.items():
                    hostname, env, repo_dir, build_cmd = val
                    shared_data[key] = {
                        "hostname":  hostname,
                        "env":       env,
                        "repo_dir":  repo_dir,
                        "build_cmd": build_cmd,
                    }
                with open(shared_path, "w") as f:
                    json.dump(shared_data, f, indent=2)
                self.log(f"✓ Tester registry synced to {shared_path}")
            except Exception as e:
                self.log(f"[WARN] Could not sync to shared folder: {e}")
                messagebox.showwarning(
                    "Shared Folder Sync Failed",
                    "Tester saved locally but could not write to shared folder:\n\n"
                    + str(e) + "\n\n"
                    "The watcher will not see this tester until the shared\n"
                    "folder is accessible. Re-save via Remove + Add Tester."
                )
        except Exception as e:
            self.log("[WARN] Could not save tester registry: " + str(e))

    def _refresh_tester_dropdown(self):
        """Rebuild the listbox values from the current registry."""
        keys = list(self._TESTER_REGISTRY.keys())
        
        # Remember current selections
        current_selections = [self._tester_listbox.get(i) for i in self._tester_listbox.curselection()]
        
        # Update listbox
        self._tester_listbox.delete(0, tk.END)
        for item in keys:
            self._tester_listbox.insert(tk.END, item)
        
        # Restore selections or select first item
        restored = False
        for i, item in enumerate(keys):
            if item in current_selections:
                self._tester_listbox.selection_set(i)
                restored = True
        
        if not restored and keys:
            self._tester_listbox.selection_set(0)
        
        self._refresh_tester_mode()

    def _refresh_tester_mode(self):
        """Update the mode badge whenever selection changes."""
        selections = self._tester_listbox.curselection()
        if not selections:
            self._compile_mode_var.set("No tester selected")
            self._compile_mode_lbl.config(foreground="#cc0000")
        elif len(selections) == 1:
            idx = selections[0]
            key = self._tester_listbox.get(idx)
            if key in self._TESTER_REGISTRY:
                val = self._TESTER_REGISTRY[key]
                hostname = val[0]
                env = val[1]
                self._compile_mode_var.set(
                    "Selected: " + hostname + "  |  Env: " + env
                )
                self._compile_mode_lbl.config(foreground="#1a6e1a")
        else:
            # Multiple testers selected
            tester_names = []
            for idx in selections:
                key = self._tester_listbox.get(idx)
                if key in self._TESTER_REGISTRY:
                    val = self._TESTER_REGISTRY[key]
                    hostname = val[0]
                    env = val[1]
                    tester_names.append(f"{hostname} ({env})")
            self._compile_mode_var.set(
                f"Multi-compile: {len(selections)} testers - " + ", ".join(tester_names)
            )
            self._compile_mode_lbl.config(foreground="#0066cc")

    def _resolve_tester(self):
        """
        Returns list of (hostname, env) tuples — one per selected tester.
        Raises ValueError with clear message if nothing is selected.
        """
        selections = self._tester_listbox.curselection()
        if not selections:
            raise ValueError(
                "No tester selected.\n\n"
                "Please select one or more testers from the list, or click\n"
                "'+ Add Tester' to register a new one."
            )
        
        targets = []
        for idx in selections:
            key = self._tester_listbox.get(idx)
            if key not in self._TESTER_REGISTRY:
                raise ValueError(
                    "Tester not found in registry: '" + key + "'\n"
                    "Please re-select from the list."
                )
            val = self._TESTER_REGISTRY[key]
            # Return (hostname, env) tuple for backward compatibility
            targets.append((val[0], val[1]))
        
        return targets

    def _open_add_tester_dialog(self):
        """
        Modal dialog to register a new tester.
        Uses a preflight checklist to ensure watcher is set up before registration.
        """
        import webbrowser

        dialog = tk.Toplevel(self.root)
        dialog.title("Add New Tester")
        dialog.geometry("560x580")
        dialog.resizable(False, False)
        dialog.grab_set()   # modal
        self._centre_dialog(dialog, 560, 580)

        pad = {"padx": 12, "pady": 6}

        # ── Header ──
        ttk.Label(dialog, text="Register New Tester",
                  font=("Arial", 12, "bold")).grid(
            row=0, column=0, columnspan=2, pady=(14, 4), **{"padx": 12})

        ttk.Separator(dialog, orient="horizontal").grid(
            row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), padx=12, pady=4)

        # ── Preflight Checklist (FIRST, not last) ──
        checklist_frame = ttk.LabelFrame(dialog, text="⚠ Preflight Checklist", padding=10)
        checklist_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), padx=12, pady=8)

        ttk.Label(checklist_frame,
                  text="Before registering, confirm the tester is ready to receive compile jobs:",
                  font=("Arial", 9, "bold"), foreground="#cc6600").pack(anchor=tk.W, pady=(0, 8))

        # Three checkboxes that must be ticked
        check_vars = []
        
        check1 = tk.BooleanVar(value=False)
        check_vars.append(check1)
        ttk.Checkbutton(checklist_frame, text="Watcher files visible at P:\\temp\\BENTO\\watcher\\",
                        variable=check1).pack(anchor=tk.W, pady=2)
        
        check2 = tk.BooleanVar(value=False)
        check_vars.append(check2)
        ttk.Checkbutton(checklist_frame, text="Watcher running (Task Scheduler configured on tester)",
                        variable=check2).pack(anchor=tk.W, pady=2)
        
        check3 = tk.BooleanVar(value=False)
        check_vars.append(check3)
        ttk.Checkbutton(checklist_frame, text="Shared folder P:\\temp\\BENTO accessible from tester",
                        variable=check3).pack(anchor=tk.W, pady=2)

        # Setup guide button
        guide_path = os.path.join(
            os.path.dirname(os.path.abspath(__file__)),
            "BENTO_Watcher_Setup_Guide.pdf"
        )
        def open_guide():
            if os.path.exists(guide_path):
                try:
                    os.startfile(guide_path)
                except Exception:
                    webbrowser.open("file:///" + guide_path.replace("\\", "/"))
            else:
                messagebox.showwarning(
                    "Guide Not Found",
                    "Setup guide PDF not found at:\n" + guide_path + "\n\n"
                    "Ask your team lead for BENTO_Watcher_Setup_Guide.pdf"
                )
        
        ttk.Button(checklist_frame, text="📄 Open Setup Guide (PDF)",
                   command=open_guide).pack(anchor=tk.W, pady=(8, 0))

        ttk.Separator(dialog, orient="horizontal").grid(
            row=3, column=0, columnspan=2, sticky=(tk.W, tk.E), padx=12, pady=8)

        # ── Tester Details (SECOND, after checklist) ──
        ttk.Label(dialog, text="Tester Hostname:").grid(
            row=4, column=0, sticky=tk.W, **pad)
        hostname_var = tk.StringVar()
        hostname_entry = ttk.Entry(dialog, textvariable=hostname_var, width=28)
        hostname_entry.grid(row=4, column=1, sticky=tk.W, **pad)
        ttk.Label(dialog, text="e.g.  IBIR-0999",
                  font=("Arial", 8), foreground="gray").grid(
            row=5, column=1, sticky=tk.W, padx=12, pady=0)

        ttk.Label(dialog, text="Environment:").grid(
            row=6, column=0, sticky=tk.W, **pad)
        # Derive ENV list from registry for consistency
        from watcher.watcher_config import TESTER_REGISTRY
        env_values = list(TESTER_REGISTRY.keys()) if TESTER_REGISTRY else ["ABIT", "SFN2", "CNFG"]
        env_var = tk.StringVar(value=env_values[0] if env_values else "")
        env_combo = ttk.Combobox(
            dialog, textvariable=env_var,
            values=env_values,
            state="readonly", width=10)
        env_combo.grid(row=6, column=1, sticky=tk.W, **pad)

        ttk.Label(dialog, text="Repo Path:").grid(
            row=7, column=0, sticky=tk.W, **pad)
        repo_dir_var = tk.StringVar(value=r"C:\BENTO\adv_ibir_master")
        repo_frame = ttk.Frame(dialog)
        repo_frame.grid(row=7, column=1, sticky=tk.W, **pad)
        ttk.Entry(repo_frame, textvariable=repo_dir_var, width=32).pack(side=tk.LEFT)
        ttk.Button(repo_frame, text="📁", width=3,
                   command=lambda: self._browse_directory(repo_dir_var, "Select TP Repository")).pack(side=tk.LEFT, padx=(2, 0))

        ttk.Label(dialog, text="Build Command:").grid(
            row=8, column=0, sticky=tk.W, **pad)
        build_cmd_var = tk.StringVar(value="make release")
        build_combo = ttk.Combobox(
            dialog, textvariable=build_cmd_var,
            values=["make release", "make release_supermicro"],
            width=25)
        build_combo.grid(row=8, column=1, sticky=tk.W, **pad)

        ttk.Separator(dialog, orient="horizontal").grid(
            row=9, column=0, columnspan=2, sticky=(tk.W, tk.E), padx=12, pady=8)

        # ── Error label ──
        err_var = tk.StringVar(value="")
        ttk.Label(dialog, textvariable=err_var,
                  foreground="red", font=("Arial", 8)).grid(
            row=10, column=0, columnspan=2, padx=12)

        # ── Buttons ──
        def _confirm():
            hostname = hostname_var.get().strip().upper()
            env      = env_var.get().strip().upper()
            repo_dir = repo_dir_var.get().strip()
            build_cmd = build_cmd_var.get().strip()
            
            if not hostname:
                err_var.set("Hostname cannot be empty.")
                return
            if not env:
                err_var.set("Please select an environment.")
                return
            if not repo_dir:
                err_var.set("Repo path cannot be empty.")
                return
            if not build_cmd:
                err_var.set("Build command cannot be empty.")
                return
            
            key = hostname + " (" + env + ")"
            if key in self._TESTER_REGISTRY:
                err_var.set("Tester '" + key + "' is already registered.")
                return
            
            self._TESTER_REGISTRY[key] = (hostname, env, repo_dir, build_cmd)
            self._save_tester_registry()
            self._refresh_tester_dropdown()
            # Select the new tester in listbox
            for i in range(self._tester_listbox.size()):
                if self._tester_listbox.get(i) == key:
                    self._tester_listbox.selection_clear(0, tk.END)
                    self._tester_listbox.selection_set(i)
                    self._tester_listbox.see(i)
                    break
            self._refresh_tester_mode()
            self.log("[Tester Added] " + key + " registered and selected.")
            dialog.destroy()

        btn_row = ttk.Frame(dialog)
        btn_row.grid(row=11, column=0, columnspan=2, pady=10)
        
        add_btn = ttk.Button(btn_row, text="Add Tester", command=_confirm, state="disabled")
        add_btn.pack(side=tk.LEFT, padx=6)
        ttk.Button(btn_row, text="Cancel", command=dialog.destroy).pack(side=tk.LEFT, padx=6)

        # ── Enable "Add Tester" only when all checkboxes are ticked ──
        def _update_add_btn(*_):
            all_checked = all(v.get() for v in check_vars)
            add_btn.config(state="normal" if all_checked else "disabled")
            if all_checked:
                err_var.set("")  # Clear any previous errors when checklist complete
        
        for v in check_vars:
            v.trace_add("write", _update_add_btn)

        hostname_entry.focus_set()

    def _remove_selected_tester(self):
        """Remove selected tester(s) from registry"""
        selections = self._tester_listbox.curselection()
        if not selections:
            messagebox.showwarning("No Selection", "Please select one or more testers to remove.")
            return
        
        # Get all selected tester names
        testers_to_remove = [self._tester_listbox.get(idx) for idx in selections]
        
        if len(testers_to_remove) == 1:
            msg = f"Remove '{testers_to_remove[0]}' from the registry?"
        else:
            msg = f"Remove {len(testers_to_remove)} testers from the registry?\n\n" + "\n".join(f"  • {t}" for t in testers_to_remove)
        
        if messagebox.askyesno("Remove Tester(s)", msg):
            for key in testers_to_remove:
                if key in self._TESTER_REGISTRY:
                    del self._TESTER_REGISTRY[key]
                    self.log(f"[Tester Removed] {key}")
            
            self._save_tester_registry()
            self._refresh_tester_dropdown()

    def _browse_directory(self, var, title):
        """Open directory browser and set the variable"""
        from tkinter import filedialog
        path = filedialog.askdirectory(title=title)
        if path:
            var.set(path)

    def _centre_dialog(self, dialog, w, h):
        """Position dialog at the centre of the main window."""
        dialog.transient(self.root)  # Keep dialog above parent
        dialog.update_idletasks()    # Force Tk to compute dialog size
        x = self.root.winfo_x() + (self.root.winfo_width()  - w) // 2
        y = self.root.winfo_y() + (self.root.winfo_height() - h) // 2
        dialog.geometry(f"{w}x{h}+{x}+{y}")

    def trigger_compile_with_lock(self):
        """Lock GUI and prompt for custom labels if multi-compile, then run in background thread."""
        try:
            targets = self._resolve_tester()
        except ValueError as e:
            messagebox.showerror("Configuration Error", str(e))
            return
            
        if len(targets) > 1:
            self._open_multi_label_dialog(targets)
        else:
            default_label = self.tgz_label_var.get().strip()
            self._start_compile_thread([(targets[0][0], targets[0][1], default_label)])

    def _open_multi_label_dialog(self, targets):
        dialog = tk.Toplevel(self.root)
        dialog.title("Set Custom TGZ Labels")
        dialog.geometry("500x" + str(180 + len(targets) * 40))
        dialog.grab_set()
        self._centre_dialog(dialog, 500, 180 + len(targets) * 40)
        
        ttk.Label(dialog, text="Set Custom TGZ Labels", style="Subtitle.TLabel").pack(pady=(15, 5))
        ttk.Label(dialog, text="Enter a specific TGZ label for each tester (leave blank for no label):").pack(pady=(0, 10))
        
        frame = ttk.Frame(dialog, padding=10)
        frame.pack(fill=tk.BOTH, expand=True, padx=20)
        
        label_vars = {}
        default_label = self.tgz_label_var.get().strip()
        
        for i, (hostname, env) in enumerate(targets):
            ttk.Label(frame, text=f"{hostname} ({env}):", width=25, font=('Segoe UI', 10, 'bold')).grid(row=i, column=0, pady=5, sticky=tk.W)
            var = tk.StringVar(value=default_label)
            label_vars[(hostname, env)] = var
            ttk.Entry(frame, textvariable=var, width=30).grid(row=i, column=1, pady=5, sticky=tk.W)
            
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(pady=15)
        
        def on_confirm():
            targets_with_labels = []
            for (hostname, env) in targets:
                targets_with_labels.append((hostname, env, label_vars[(hostname, env)].get().strip()))
            dialog.destroy()
            self._start_compile_thread(targets_with_labels)
            
        ttk.Button(btn_frame, text="Confirm & Compile", style="Compile.TButton", command=on_confirm).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Cancel", command=dialog.destroy).pack(side=tk.LEFT, padx=5)

    def _start_compile_thread(self, targets_with_labels):
        import threading

        # ── Resolve tester on main thread (tkinter widget access) ──
        try:
            targets = self._resolve_tester()
        except ValueError as e:
            messagebox.showerror("No Tester Selected", str(e))
            return

        # ── Collect all widget values on main thread ──
        issue_key = self.impl_issue_var.get().strip().upper()
        repo_path = self.impl_repo_var.get().strip()
        raw_zip   = self.raw_zip_var.get().strip()
        release   = self.release_tgz_var.get().strip()
        label     = self.tgz_label_var.get().strip()

        self.lock_gui()
        self.compile_status_var.set("Compiling...")
        self.compile_btn.config(state="disabled")

        # Save config on main thread before thread starts (widget access not allowed in thread)
        self.save_config()

        # Open live status monitor for single-tester compiles
        if len(targets) == 1:
            hostname, env = targets[0]
            if issue_key and raw_zip:
                self.root.after(100, lambda: self.open_build_status_monitor(
                    issue_key, hostname, env, raw_zip))

        t = threading.Thread(
            target=self._run_compile_thread,
            args=(targets, issue_key, repo_path, raw_zip, release, label),
            daemon=True
        )
        t.start()

    def _run_compile_thread(self, targets, issue_key, repo_path, raw_zip, release, label):
        """Background thread: runs compile and updates status on completion."""
        try:
            self._run_compile(targets, issue_key, repo_path, raw_zip, release, label)
        except Exception as e:
            self.log(f"[Compile error] {str(e)}")
            self.compile_status_var.set("Error - check log")
        finally:
            self.unlock_gui()
            self.compile_btn.config(state="normal")

    def _run_compile(self, targets, issue_key, repo_path, raw_zip, release, label):
        """Main compile logic called from background thread.
        All parameters pre-resolved on main thread — no widget access here.
        """
        import os

        # ── Pre-flight checks ──
        errors = []
        if not issue_key:
            errors.append("JIRA Issue Key is required")
        if not repo_path:
            errors.append("Local Repo Path is required")
        elif not os.path.isdir(repo_path):
            errors.append("Repo path not found: " + repo_path)
        if not raw_zip:
            errors.append("RAW_ZIP Path is required")
        elif not os.path.isdir(raw_zip):
            errors.append("RAW_ZIP folder not accessible: " + raw_zip +
                         "\n    Check that the P: drive is mapped on this machine")
        if not release:
            errors.append("RELEASE_TGZ Path is required")
        elif not os.path.isdir(release):
            errors.append("RELEASE_TGZ folder not accessible: " + release +
                         "\n    Check that the P: drive is mapped on this machine")

        if not targets:
            errors.append("No tester selected.")

        if errors:
            msg = "Cannot compile. Fix the following:\n\n" + "\n".join(
                "  " + chr(8226) + " " + e for e in errors)
            messagebox.showerror("Configuration Error", msg)
            self.compile_status_var.set("Config error - not compiled")
            return

        # ── Load orchestrator ──
        try:
            import compilation_orchestrator as orch
        except ImportError:
            messagebox.showerror(
                "Missing File",
                "compilation_orchestrator.py not found next to main.py."
            )
            self.compile_status_var.set("Missing orchestrator")
            return

        # ── Wrap log callback to detect phase updates ──
        def _compile_log(msg):
            if msg.startswith("__PHASE__:"):
                phase_msg = msg[10:]
                # Truncate to fit status label
                if len(phase_msg) > 45:
                    phase_msg = phase_msg[:42] + "..."
                # Update status label from main thread
                self.root.after(0, lambda: self.compile_status_var.set(phase_msg))
            else:
                self.log(msg)

        sep = "=" * 60
        self.log("\n" + sep)
        self.log("  COMPILE TP PACKAGE")
        self.log("  JIRA Issue : " + issue_key)
        
        # Handle single or multiple testers
        if len(targets_with_labels) == 1:
            hostname, env, label = targets_with_labels[0]
            self.log("  Tester     : " + hostname + " (" + env + ")")
            self.log("  Label      : " + (label if label else "(none)"))
        else:
            self.log(f"  Testers    : {len(targets_with_labels)} parallel compilations")
            for hostname, env, label in targets_with_labels:
                self.log(f"               - {hostname} ({env}) [Label: {label if label else '(none)'}]")
        
        self.log("  Repo Path  : " + repo_path)
        self.log("  RAW_ZIP    : " + raw_zip)
        self.log("  RELEASE    : " + release)
        self.log(sep)

        # Use multi-tester compilation if multiple targets selected
        if len(targets_with_labels) == 1:
            # Single tester - use original flow
            hostname, env, label = targets_with_labels[0]
            result = orch.compile_tp_package(
                source_dir=repo_path,
                env=env,
                jira_key=issue_key,
                hostname=hostname,
                label=label,
                raw_zip_folder=raw_zip,
                release_tgz_folder=release,
                log_callback=_compile_log,
            )
            self._handle_single_compile_result(result)
        else:
            # Multiple testers — open live status window first
            callbacks, status_win = self.open_multi_compile_status(targets)

            # Mark all as building before threads start
            for hostname, env in targets:
                self.root.after(0, lambda h=hostname, e=env:
                    callbacks[(h, e)]("building"))

            # Wrap compile_tp_package_multi to update badges as each completes
            def _one_with_badge(hostname, env, label):
                self.root.after(0, lambda: callbacks[(hostname, env)]("building"))
                result = orch.compile_tp_package(
                    source_dir=repo_path,
                    env=env,
                    jira_key=issue_key,
                    hostname=hostname,
                    label=label,
                    raw_zip_folder=raw_zip,
                    release_tgz_folder=release,
                    log_callback=_compile_log,
                )
                result["hostname"] = hostname
                result["env"]      = env
                state = result.get("status", "failed")
                elapsed = result.get("elapsed", None)
                self.root.after(0, lambda s=state, el=elapsed:
                    callbacks[(hostname, env)](s, el))
                return result

            from concurrent.futures import ThreadPoolExecutor, as_completed
            results = []
            with ThreadPoolExecutor(max_workers=len(targets_with_labels)) as pool:
                futures = {pool.submit(_one_with_badge, h, e, l): (h, e)
                           for h, e, l in targets_with_labels}
                for future in as_completed(futures):
                    results.append(future.result())

            self._handle_multi_compile_results(results)

    def _build_compile_history_section(self, parent_tab):
        """
        Compile History section — reads build_info_*.txt files from
        RELEASE_TGZ and displays them in a sortable table.
        Placed inside the Implementation tab.
        """
        import glob

        history_frame = ttk.LabelFrame(parent_tab, text="📋 Compile History", padding=8)
        history_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        # Toolbar
        toolbar = ttk.Frame(history_frame)
        toolbar.pack(fill=tk.X, pady=(0, 5))
        ttk.Label(toolbar, text="Past compilations from RELEASE_TGZ",
                  font=("Arial", 9), foreground="gray").pack(side=tk.LEFT)
        ttk.Button(toolbar, text="🔄 Refresh", command=lambda: _load()).pack(side=tk.RIGHT)
        ttk.Button(toolbar, text="📁 Open Folder",
                   command=lambda: os.startfile(
                       self.release_tgz_var.get().strip()
                       if hasattr(self, "release_tgz_var")
                       else r"P:\temp\BENTO\RELEASE_TGZ"
                   )).pack(side=tk.RIGHT, padx=5)

        # Table
        cols = ("Timestamp", "JIRA", "Tester", "ENV", "Label", "Output TGZ")
        tree = ttk.Treeview(history_frame, columns=cols, show="headings", height=8)
        col_widths = [140, 120, 110, 70, 100, 220]
        for col, w in zip(cols, col_widths):
            tree.heading(col, text=col, command=lambda c=col: _sort(c))
            tree.column(col, width=w, anchor="w")

        tree.tag_configure("even", background="#f5f5f5")
        tree.tag_configure("odd",  background="#ffffff")

        scroll_y = ttk.Scrollbar(history_frame, orient=tk.VERTICAL, command=tree.yview)
        tree.configure(yscrollcommand=scroll_y.set)
        tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scroll_y.pack(side=tk.LEFT, fill=tk.Y)

        def _load():
            release_root = (self.release_tgz_var.get().strip()
                            if hasattr(self, "release_tgz_var")
                            else r"P:\temp\BENTO\RELEASE_TGZ")

            def _fetch():
                pattern = os.path.join(release_root, "**", "build_info_*.txt")
                files   = glob.glob(pattern, recursive=True)
                files.sort(key=os.path.getmtime, reverse=True)
                self.root.after(0, lambda: _populate(files))

            def _populate(files):
                tree.delete(*tree.get_children())
                for i, fpath in enumerate(files):
                    row = _parse(fpath)
                    if row:
                        tree.insert("", tk.END, values=row,
                                    tags=("even" if i % 2 == 0 else "odd",))

            threading.Thread(target=_fetch, daemon=True).start()

        def _parse(fpath):
            data = {}
            try:
                with open(fpath, "r") as f:
                    for line in f:
                        if ":" in line:
                            k, _, v = line.partition(":")
                            data[k.strip()] = v.strip()
                return (data.get("Timestamp", ""), data.get("JIRA Key", ""),
                        data.get("Tester", ""),    data.get("Env", ""),
                        data.get("Label", ""),     data.get("Output TGZ", ""))
            except Exception:
                return None

        def _sort(col):
            items = [(tree.set(k, col), k) for k in tree.get_children("")]
            items.sort()
            for idx, (_, k) in enumerate(items):
                tree.move(k, "", idx)

        _load()

    def _handle_single_compile_result(self, result):
        """Handle result from single-tester compilation"""
        status   = result.get("status",   "failed")
        detail   = result.get("detail",   "Unknown error")
        elapsed  = result.get("elapsed",  0)
        tgz      = result.get("tgz_file", "")
        tgz_path = result.get("tgz_path", "")
        zip_file = result.get("zip_file", "")

        # ── Open live log viewer if we can locate the bento_status sidecar ──
        if zip_file:
            raw_zip = self.raw_zip_var.get().strip() if hasattr(self, "raw_zip_var") else ""
            status_path = os.path.join(raw_zip, zip_file + ".bento_status") if raw_zip else ""
            build_log_hint = "C:\\BENTO\\logs\\build_" + os.path.splitext(zip_file)[0] + ".log"
        else:
            status_path    = ""
            build_log_hint = ""

        if status == "success":
            self.compile_status_var.set("Done: " + tgz)
            self.log("\n[OK] COMPILE SUCCESS")
            self.log("  TGZ file : " + tgz)
            self.log("  Location : " + tgz_path)
            self.log("  Time     : " + str(int(elapsed)) + "s")
            if build_log_hint:
                self.log("  Build log: " + build_log_hint + "  (on tester machine)")
            messagebox.showinfo(
                "Compile Success",
                "TGZ ready:\n" + tgz_path
            )

        elif status == "timeout":
            self.compile_status_var.set("TIMEOUT")
            if build_log_hint:
                detail += "\n\nBuild log on tester:\n" + build_log_hint
            self._show_compile_error_dialog("Compile Timeout", detail, status)

        else:
            self.compile_status_var.set("FAILED")
            if build_log_hint:
                detail += "\n\nBuild log on tester:\n" + build_log_hint
            self._show_compile_error_dialog("Compile Failed", detail, status)

    def _handle_multi_compile_results(self, results):
        """Handle results from multi-tester parallel compilation"""
        success_count = sum(1 for r in results if r.get("status") == "success")
        failed_count = sum(1 for r in results if r.get("status") == "failed")
        timeout_count = sum(1 for r in results if r.get("status") == "timeout")
        
        self.log("\n" + "=" * 60)
        self.log("  MULTI-TESTER COMPILATION RESULTS")
        self.log("=" * 60)
        
        tgz_paths = []
        for r in results:
            hostname = r.get("hostname", "Unknown")
            env = r.get("env", "Unknown")
            status = r.get("status", "failed")
            tag = f"{hostname} ({env})"
            
            if status == "success":
                tgz_file = r.get("tgz_file", "")
                tgz_path = r.get("tgz_path", "")
                elapsed = r.get("elapsed", 0)
                self.log(f"[OK] {tag}")
                self.log(f"     TGZ: {tgz_file}")
                self.log(f"     Time: {int(elapsed)}s")
                tgz_paths.append(tgz_path)
            elif status == "timeout":
                detail   = r.get("detail", "Timeout")
                zip_file = r.get("zip_file", "")
                self.log(f"[TIMEOUT] {tag}")
                self.log(f"          {detail}")
                if zip_file:
                    self.log("          Build log on tester: C:\\BENTO\\logs\\build_"
                             + os.path.splitext(zip_file)[0] + ".log")
            else:
                detail   = r.get("detail", "Unknown error")
                zip_file = r.get("zip_file", "")
                self.log(f"[FAIL] {tag}")
                self.log(f"       {detail}")
                if zip_file:
                    self.log("       Build log on tester: C:\\BENTO\\logs\\build_"
                             + os.path.splitext(zip_file)[0] + ".log")
        
        self.log("=" * 60)
        self.log(f"Summary: {success_count} success, {failed_count} failed, {timeout_count} timeout")
        self.log("=" * 60)
        
        # Update status and show summary dialog
        if success_count == len(results):
            self.compile_status_var.set(f"All {len(results)} compilations succeeded")
            messagebox.showinfo(
                "Multi-Compile Success",
                f"All {len(results)} testers compiled successfully!\n\n" +
                "\n".join(f"✓ {os.path.basename(p)}" for p in tgz_paths)
            )
        elif success_count > 0:
            self.compile_status_var.set(f"{success_count}/{len(results)} succeeded")
            messagebox.showwarning(
                "Partial Success",
                f"{success_count} of {len(results)} compilations succeeded.\n"
                f"{failed_count} failed, {timeout_count} timed out.\n\n"
                "Check the log for details."
            )
        else:
            self.compile_status_var.set("All compilations failed")
            messagebox.showerror(
                "Multi-Compile Failed",
                f"All {len(results)} compilations failed.\n\n"
                "Check the log for details."
            )

    def _show_compile_error_dialog(self, title, detail, status):
        """
        Rich error dialog for compile failures.
        Shows full error detail with scroll, and actionable hints.
        """
        dialog = tk.Toplevel(self.root)
        dialog.title(title)
        dialog.geometry("600x420")
        dialog.grab_set()
        self._centre_dialog(dialog, 600, 420)

        # Header
        color = "#cc0000" if status == "failed" else "#b36200"
        icon  = "[FAILED]" if status == "failed" else "[TIMEOUT]"
        ttk.Label(dialog, text=icon + "  " + title,
                  font=("Arial", 12, "bold"), foreground=color).pack(
            anchor=tk.W, padx=16, pady=(14, 4))

        ttk.Separator(dialog, orient="horizontal").pack(
            fill=tk.X, padx=16, pady=4)

        ttk.Label(dialog, text="Error detail from tester:",
                  font=("Arial", 9, "bold")).pack(anchor=tk.W, padx=16, pady=(4, 2))

        # Scrollable error text
        err_frame = ttk.Frame(dialog)
        err_frame.pack(fill=tk.BOTH, expand=True, padx=16, pady=4)
        scrollbar = ttk.Scrollbar(err_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        err_text = tk.Text(err_frame, height=10, wrap=tk.WORD,
                           yscrollcommand=scrollbar.set,
                           font=("Courier", 9), foreground="#333333",
                           background="#fff8f8")
        err_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=err_text.yview)
        err_text.insert("1.0", detail)
        err_text.config(state="disabled")

        ttk.Separator(dialog, orient="horizontal").pack(
            fill=tk.X, padx=16, pady=4)

        ttk.Label(dialog, text="Common causes:",
                  font=("Arial", 9, "bold")).pack(anchor=tk.W, padx=16)
        hints = (
            "  - Watcher script not running on tester  (check Task Scheduler)\n"
            "  - Wrong tester selected  (ZIP env token must match watcher --env)\n"
            "  - RAW_ZIP or RELEASE_TGZ path not accessible from tester\n"
            "  - make release compilation error  (check build log on tester)\n"
            "  - Tester out of disk space or RAM"
        )
        ttk.Label(dialog, text=hints, justify=tk.LEFT,
                  font=("Arial", 8), foreground="#555555").pack(
            anchor=tk.W, padx=16, pady=(2, 8))

        # Buttons: Retry + Close
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(pady=(0, 12))
        
        def _retry():
            dialog.destroy()
            self.trigger_compile_with_lock()
        
        ttk.Button(btn_frame, text="🔄 Retry Compile", command=_retry).pack(side=tk.LEFT, padx=4)
        ttk.Button(btn_frame, text="Close", command=dialog.destroy).pack(side=tk.LEFT, padx=4)

    def generate_implementation_only_with_lock(self):
        """Wrapper for generate_implementation_only with GUI locking"""
        self.lock_gui()
        try:
            return self.generate_implementation_only()
        finally:
            self.unlock_gui()
    
    def generate_implementation_only(self):
        """Generate implementation plan only"""
        repo_path = self.impl_repo_var.get().strip()
        if not repo_path:
            messagebox.showerror("Error", "Please enter repository path")
            return
        
        # Get issue key from the implementation tab field (auto-populated from home)
        issue_key = self.impl_issue_var.get().strip().upper()
        if not issue_key or issue_key == f"{self.jira_project_var.get()}-":
            messagebox.showerror("Error", "Please enter JIRA issue key")
            return
        
        self.log(f"\n[Generating Implementation Plan for {issue_key}]")
        
        # Initialize workflow file to check for existing data
        self.init_workflow_file(issue_key)
        
        # Check if we already have an implementation plan in workflow
        existing_impl_plan = self.get_workflow_step("IMPLEMENTATION_PLAN")
        if existing_impl_plan:
            self.log(f"✓ Found existing implementation plan in workflow file")
            self.log(f"  Opening interactive chat with existing plan...")
            
            # Extract just the plan content (remove header)
            plan_lines = existing_impl_plan.split('\n')
            # Skip the header lines (# Implementation Plan, **JIRA Issue:**, **Generated:**)
            plan_content = '\n'.join([line for line in plan_lines if not line.startswith('#') and not line.startswith('**')])
            
            # Open chat with existing plan
            self.root.after(100, lambda: self.create_interactive_chat(
                issue_key,
                "Implementation Plan",
                plan_content.strip(),
                lambda: self.finalize_implementation_plan(repo_path, issue_key, plan_content.strip())
            ))
            return
        
        # Index repository
        repo_index = self.analyzer.index_repository(repo_path)
        
        # Check for existing JIRA analysis in workflow
        jira_analysis_text = self.get_workflow_step("JIRA_ANALYSIS")
        if jira_analysis_text:
            self.log("✓ Using existing JIRA analysis from workflow file")
            jira_analysis = {'success': True, 'analysis': jira_analysis_text}
        else:
            self.log("⚠ JIRA analysis not found in workflow, fetching and analyzing...")
            issue_data = self.analyzer.fetch_jira_issue(issue_key)
            if not issue_data:
                messagebox.showerror("Error", f"Failed to fetch JIRA issue {issue_key}")
                return
            
            jira_analysis = self.analyzer.analyze_jira_request(issue_data)
            if not jira_analysis.get('success'):
                messagebox.showerror("Error", f"Failed to analyze JIRA issue: {jira_analysis.get('error')}")
                return
            
            # Save to workflow for future use
            self.save_workflow_step("JIRA_ANALYSIS", jira_analysis.get('analysis', ''))
        
        # Check for existing impact analysis in workflow
        impact_analysis_text = self.get_workflow_step("IMPACT_ANALYSIS")
        if impact_analysis_text:
            self.log("✓ Using existing impact analysis from workflow file")
            impact_analysis = {'success': True, 'impact_analysis': impact_analysis_text}
        else:
            self.log("⚠ Impact analysis not found in workflow, generating...")
            impact_analysis = self.analyzer.analyze_code_impact(repo_path, repo_index, jira_analysis)
            if not impact_analysis.get('success'):
                messagebox.showerror("Error", f"Failed to analyze code impact: {impact_analysis.get('error')}")
                return
            
            # Save to workflow for future use
            self.save_workflow_step("IMPACT_ANALYSIS", impact_analysis.get('impact_analysis', ''))
        
        # Generate implementation plan with interactive chat
        self.generate_implementation_with_chat(repo_path, issue_key, jira_analysis, impact_analysis, repo_index)
    
    def generate_implementation_with_chat(self, repo_path, issue_key, jira_analysis, impact_analysis, repo_index):
        """Generate implementation plan and open interactive chat for review"""
        import os
        
        self.log("\n[Generating Implementation Plan with AI]")
        
        # Get list of key files to analyze
        key_files = []
        for file_info in repo_index['file_index'][:20]:  # Top 20 files
            file_path = os.path.join(repo_path, file_info['path'])
            if os.path.isfile(file_path) and file_info['size'] < 100000:  # Skip large files
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                        key_files.append({
                            'path': file_info['path'],
                            'content': content[:2000]  # First 2000 chars
                        })
                except:
                    pass
        
        # Prepare implementation prompt
        impl_prompt = f"""Based on the JIRA analysis and code impact assessment, provide specific code changes:

**JIRA Analysis:**
{jira_analysis.get('analysis', 'N/A')}

**Impact Analysis:**
{impact_analysis.get('impact_analysis', 'N/A')}

**Sample Code Files:**
{json.dumps(key_files[:5], indent=2)}

Provide implementation in this EXACT format:

## FILES_TO_MODIFY
- path/to/file1.py: Brief description of changes
- path/to/file2.js: Brief description of changes

## CODE_CHANGES
### path/to/file1.py
```python
# Complete updated code for this file
```

### path/to/file2.js
```javascript
// Complete updated code for this file
```

## NEW_FILES
### path/to/newfile.py
```python
# Complete code for new file
```

## IMPLEMENTATION_GUIDE
Step-by-step guide for manual review and additional changes.
"""
        
        messages = [{"role": "user", "content": impl_prompt}]
        
        # Use code generation model
        result = self.analyzer.ai_client.chat_completion(messages, task_type="code_generation", mode="code_implementation")
        
        if result['success']:
            implementation_plan = result['response']['choices'][0]['message']['content']
            self.log("✓ AI generated implementation plan")
            
            # Open interactive chat for review before saving
            self.log("✓ Opening interactive chat for implementation plan review...")
            self.root.after(100, lambda: self.create_interactive_chat(
                issue_key,
                "Implementation Plan",
                implementation_plan,
                lambda: self.finalize_implementation_plan(repo_path, issue_key, implementation_plan)
            ))
        else:
            self.log(f"✗ Failed to generate implementation plan: {result['error']}")
            messagebox.showerror("Error", f"Failed to generate implementation plan:\n{result['error']}")
    
    def finalize_implementation_plan(self, repo_path, issue_key, implementation_plan):
        """Finalize implementation plan after user approval from interactive chat"""
        import os
        
        # Save implementation plan to workflow file instead of separate .md file
        formatted_plan = f"# Implementation Plan\n\n"
        formatted_plan += f"**JIRA Issue:** {issue_key}\n"
        formatted_plan += f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        formatted_plan += implementation_plan
        
        self.save_workflow_step("IMPLEMENTATION_PLAN", formatted_plan, issue_key)
        self.log(f"✓ Implementation plan saved to workflow file")
        
        # Update result window
        self.impl_result_text.delete('1.0', tk.END)
        self.impl_result_text.insert('1.0', f"Implementation plan generated and saved to workflow!\n\n{implementation_plan[:500]}...")
        
        # Ask if user wants to apply changes automatically
        apply_auto = messagebox.askyesno("Apply Code Changes",
            "Implementation plan has been saved to workflow.\n\n"
            "Would you like to apply code changes automatically?\n\n"
            "This will modify files based on AI suggestions.\n"
            "Git diff will be shown after changes are applied.")
        
        if apply_auto:
            changes_applied = self.analyzer._apply_code_changes(repo_path, implementation_plan)
            if changes_applied > 0:
                self.log(f"\n✓ Applied {changes_applied} code changes")
                self.log("  Showing git diff...")
                
                # Show git diff in a window
                self.show_git_diff(repo_path)
                
                messagebox.showinfo("Success",
                    f"Applied {changes_applied} code changes!\n\n"
                    f"Git diff window has been opened for review.")
            else:
                self.log("\n  No automatic changes applied. Review workflow file for manual changes.")
                messagebox.showinfo("Manual Review Required",
                    "No automatic changes were applied.\n\n"
                    "Please review the workflow file for manual implementation.")
        else:
            self.log("\n  Review workflow file and apply changes manually.")
            messagebox.showinfo("Manual Review",
                "Implementation plan saved to workflow file.\n\n"
                "Please review and apply changes manually.")
        
        self.log(f"✓ Implementation plan for {issue_key} is complete")
    
    def show_git_diff(self, repo_path):
        """Show git diff in a new window after code changes are applied"""
        import subprocess
        
        try:
            # Get git diff
            result = subprocess.run(
                ['git', 'diff'],
                cwd=repo_path,
                capture_output=True,
                text=True,
                timeout=30
            )
            
            diff_output = result.stdout
            
            if not diff_output.strip():
                # Try git diff --cached for staged changes
                result = subprocess.run(
                    ['git', 'diff', '--cached'],
                    cwd=repo_path,
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                diff_output = result.stdout
            
            if not diff_output.strip():
                self.log("  No changes detected in git diff")
                return
            
            # Create diff window
            diff_window = tk.Toplevel(self.root)
            diff_window.title(f"Git Diff - {os.path.basename(repo_path)}")
            diff_window.geometry("1000x700")
            
            # Header
            header_frame = ttk.Frame(diff_window, padding="10")
            header_frame.pack(fill=tk.X)
            
            ttk.Label(header_frame, text="Git Diff - Code Changes",
                     font=('Arial', 12, 'bold')).pack(anchor=tk.W)
            ttk.Label(header_frame, text=f"Repository: {repo_path}",
                     font=('Arial', 9), foreground='gray').pack(anchor=tk.W)
            
            # Diff display
            diff_frame = ttk.LabelFrame(diff_window, text="Changes", padding="10")
            diff_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
            
            diff_text = scrolledtext.ScrolledText(diff_frame, height=35, width=120,
                                                  wrap=tk.NONE, font=('Courier', 9))
            diff_text.pack(fill=tk.BOTH, expand=True)
            
            # Insert diff with syntax highlighting
            for line in diff_output.split('\n'):
                if line.startswith('+') and not line.startswith('+++'):
                    diff_text.insert(tk.END, line + '\n', 'added')
                elif line.startswith('-') and not line.startswith('---'):
                    diff_text.insert(tk.END, line + '\n', 'removed')
                elif line.startswith('@@'):
                    diff_text.insert(tk.END, line + '\n', 'hunk')
                elif line.startswith('diff --git'):
                    diff_text.insert(tk.END, line + '\n', 'file')
                else:
                    diff_text.insert(tk.END, line + '\n')
            
            # Configure tags for syntax highlighting
            diff_text.tag_config('added', foreground='green')
            diff_text.tag_config('removed', foreground='red')
            diff_text.tag_config('hunk', foreground='blue', font=('Courier', 9, 'bold'))
            diff_text.tag_config('file', foreground='purple', font=('Courier', 9, 'bold'))
            
            diff_text.config(state=tk.DISABLED)
            
            # Buttons
            btn_frame = ttk.Frame(diff_window, padding="10")
            btn_frame.pack(fill=tk.X)
            
            ttk.Button(btn_frame, text="Copy to Clipboard",
                      command=lambda: self.copy_to_clipboard(diff_output)).pack(side=tk.LEFT, padx=5)
            ttk.Button(btn_frame, text="Save to File",
                      command=lambda: self.save_diff_to_file(diff_output, repo_path)).pack(side=tk.LEFT, padx=5)
            ttk.Button(btn_frame, text="Close",
                      command=diff_window.destroy).pack(side=tk.RIGHT, padx=5)
            
            self.log(f"✓ Git diff window opened")
            
        except subprocess.TimeoutExpired:
            self.log("✗ Git diff timed out")
            messagebox.showerror("Error", "Git diff command timed out")
        except Exception as e:
            self.log(f"✗ Error showing git diff: {str(e)}")
            messagebox.showerror("Error", f"Failed to show git diff:\n{str(e)}")
    
    def copy_to_clipboard(self, text):
        """Copy text to clipboard"""
        self.root.clipboard_clear()
        self.root.clipboard_append(text)
        self.log("✓ Copied to clipboard")
        messagebox.showinfo("Success", "Diff copied to clipboard!")
    
    def save_diff_to_file(self, diff_output, repo_path):
        """Save diff output to a file"""
        from tkinter import filedialog
        
        filename = filedialog.asksaveasfilename(
            title="Save Git Diff",
            defaultextension=".diff",
            filetypes=[("Diff files", "*.diff"), ("Text files", "*.txt"), ("All files", "*.*")],
            initialdir=repo_path,
            initialfile=f"changes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.diff"
        )
        
        if filename:
            try:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(diff_output)
                self.log(f"✓ Diff saved to: {filename}")
                messagebox.showinfo("Success", f"Diff saved to:\n{filename}")
            except Exception as e:
                self.log(f"✗ Error saving diff: {str(e)}")
                messagebox.showerror("Error", f"Failed to save diff:\n{str(e)}")
    
    def sync_issue_to_tabs(self, *args):
        """Sync issue key from home tab to other tabs"""
        issue_key = self.issue_var.get()
        self.fetch_issue_var.set(issue_key)
        self.analyze_issue_var.set(issue_key)
        self.repo_issue_var.set(issue_key)
        self.impl_issue_var.set(issue_key)
        self.test_issue_var.set(issue_key)
        self.risk_issue_var.set(issue_key)
    
    def sync_repo_to_tabs(self, *args):
        """Sync repository from home tab to other tabs"""
        repo = self.repo_var.get()
        self.repo_tab_var.set(repo)
    
    def sync_branch_to_tabs(self, *args):
        """Sync branch from home tab to other tabs"""
        branch = self.branch_var.get()
        self.branch_tab_var.set(branch)
    
    def toggle_debug_mode(self, *args):
        """Toggle debug mode instantly when checkbox is changed"""
        debug_enabled = self.debug_var.get()
        
        # Update AI client debug mode if it exists
        if hasattr(self.analyzer, 'ai_client') and self.analyzer.ai_client:
            if hasattr(self.analyzer.ai_client, 'config'):
                if 'settings' not in self.analyzer.ai_client.config:
                    self.analyzer.ai_client.config['settings'] = {}
                self.analyzer.ai_client.config['settings']['debug'] = debug_enabled
                self.analyzer.ai_client.debug = debug_enabled
        
        # Show/hide debug indicator
        if debug_enabled:
            self.debug_indicator.pack(side=tk.RIGHT, padx=10)
            self.log("✓ Debug mode ENABLED")
        else:
            self.debug_indicator.pack_forget()
            self.log("✓ Debug mode DISABLED")
    
    def log(self, message):
        """Thread-safe log — routes to main thread via root.after()."""
        self.root.after(0, lambda m=message: self._log_safe(m))

    def _log_safe(self, message):
        """Actual log write — must only be called from main thread."""
        self.log_text.insert(tk.END, f"{message}\n")
        self.log_text.see(tk.END)
        # Also write to log file
        if hasattr(self, 'file_logger') and self.file_logger:
            clean_msg = str(message).strip()
            if '✗' in clean_msg or 'Error' in clean_msg or 'FAILED' in clean_msg:
                self.file_logger.error(clean_msg)
            elif '⚠' in clean_msg or 'Warning' in clean_msg:
                self.file_logger.warning(clean_msg)
            elif '✓' in clean_msg or 'SUCCESS' in clean_msg:
                self.file_logger.info(clean_msg)
            elif clean_msg.startswith('===') or clean_msg.startswith('---'):
                self.file_logger.info(clean_msg)
            else:
                self.file_logger.info(clean_msg)
    
    def check_saved_credentials(self):
        if os.path.exists(CredentialManager.CREDENTIAL_FILE):
            self.log("Found saved credentials file")
    
    def lock_gui(self):
        """Lock entire GUI during operations"""
        self.gui_locked = True
        # Disable all child widgets recursively
        self._set_widget_state(self.root, 'disabled')
        self.root.config(cursor="wait")
        self.log("🔒 GUI locked during operation...")
    
    def unlock_gui(self):
        """Unlock entire GUI after operations"""
        self.gui_locked = False
        # Re-enable all child widgets recursively
        self._set_widget_state(self.root, 'normal')
        self.root.config(cursor="")
        self.log("🔓 GUI unlocked")
    
    def _set_widget_state(self, widget, state):
        """Recursively set state of all widgets except result displays"""
        try:
            # Skip all ScrolledText widgets so users can still see progress and results
            # This includes: log_text, issue_result_text, analyze_result_text,
            # repo_result_text, impact_result_text, test_result_text,
            # risk_result_text, impl_result_text
            if widget == self.log_text:
                return
            if hasattr(self, 'issue_result_text') and widget == self.issue_result_text:
                return
            if hasattr(self, 'analyze_result_text') and widget == self.analyze_result_text:
                return
            if hasattr(self, 'repo_result_text') and widget == self.repo_result_text:
                return
            if hasattr(self, 'impact_result_text') and widget == self.impact_result_text:
                return
            if hasattr(self, 'test_result_text') and widget == self.test_result_text:
                return
            if hasattr(self, 'risk_result_text') and widget == self.risk_result_text:
                return
            if hasattr(self, 'impl_result_text') and widget == self.impl_result_text:
                return
            
            # Try to set state for widgets that support it
            if hasattr(widget, 'config'):
                try:
                    widget.config(state=state)
                except:
                    pass
            
            # Recursively process children
            for child in widget.winfo_children():
                self._set_widget_state(child, state)
        except:
            pass
    
    def check_jira_connectivity(self):
        """Check if JIRA URL and token are valid"""
        if not self.analyzer.jira_token or not self.analyzer.email:
            self.log("⚠ Cannot check JIRA - credentials not loaded")
            return False
        
        self.log("\n[Checking JIRA Connectivity...]")
        
        try:
            # Try to access JIRA API
            url = f"{self.analyzer.jira_base_url}/rest/api/2/myself"
            
            auth_string = f"{self.analyzer.email}:{self.analyzer.jira_token}"
            auth_bytes = auth_string.encode('utf-8')
            auth_header = base64.b64encode(auth_bytes).decode('utf-8')
            
            headers = {
                "Authorization": f"Basic {auth_header}",
                "Content-Type": "application/json"
            }
            
            request = urllib.request.Request(url, headers=headers)
            ssl_context = ssl._create_unverified_context()
            
            response = urllib.request.urlopen(request, context=ssl_context, timeout=30)
            data = json.loads(response.read().decode('utf-8'))
            
            username = data.get('displayName', data.get('name', 'Unknown'))
            self.log(f"  ✓ JIRA Connection: SUCCESS")
            self.log(f"  ✓ Logged in as: {username}")
            self.log(f"  ✓ URL: {self.analyzer.jira_base_url}")
            return True
            
        except urllib.error.HTTPError as e:
            error_body = e.read().decode('utf-8')
            self.log(f"  ✗ JIRA Connection: FAILED (HTTP {e.code})")
            self.log(f"  ✗ Error: {error_body[:200]}")
            return False
        except Exception as e:
            self.log(f"  ✗ JIRA Connection: FAILED")
            self.log(f"  ✗ Error: {str(e)}")
            return False
    
    def check_bitbucket_connectivity(self):
        """Check if Bitbucket URL and token are valid"""
        if not self.analyzer.bitbucket_token or not self.analyzer.bitbucket_username:
            self.log("⚠ Cannot check Bitbucket - credentials not loaded")
            return False
        
        self.log("\n[Checking Bitbucket Connectivity...]")
        
        try:
            # Try to list projects
            url = f"{self.analyzer.bitbucket_base_url}/rest/api/1.0/projects?limit=1"
            
            auth_string = f"{self.analyzer.bitbucket_username}:{self.analyzer.bitbucket_token}"
            auth_bytes = auth_string.encode('utf-8')
            auth_header = base64.b64encode(auth_bytes).decode('utf-8')
            
            headers = {
                "Authorization": f"Basic {auth_header}",
                "Content-Type": "application/json"
            }
            
            request = urllib.request.Request(url, headers=headers)
            ssl_context = ssl._create_unverified_context()
            
            response = urllib.request.urlopen(request, context=ssl_context, timeout=30)
            data = json.loads(response.read().decode('utf-8'))
            
            project_count = data.get('size', 0)
            self.log(f"  ✓ Bitbucket Connection: SUCCESS")
            self.log(f"  ✓ Username: {self.analyzer.bitbucket_username}")
            self.log(f"  ✓ URL: {self.analyzer.bitbucket_base_url}")
            self.log(f"  ✓ Accessible projects: {project_count}+")
            return True
            
        except urllib.error.HTTPError as e:
            error_body = e.read().decode('utf-8')
            self.log(f"  ✗ Bitbucket Connection: FAILED (HTTP {e.code})")
            self.log(f"  ✗ Error: {error_body[:200]}")
            return False
        except Exception as e:
            self.log(f"  ✗ Bitbucket Connection: FAILED")
            self.log(f"  ✗ Error: {str(e)}")
            return False
    
    def test_model_prompt(self, model_name, model_type):
        """Test if a model can actually respond to prompts"""
        try:
            gateway_url = self.config.get('model_gateway', {}).get('base_url',
                                                                    'https://model-gateway.gcldgenaigw.gc.micron.com/api/v1')
            url = f"{gateway_url}/chat/completions"
            
            payload = {
                "model": model_name,
                "messages": [{"role": "user", "content": "Say 'OK' if you can read this."}]
            }
            
            headers = {
                "Authorization": f"Bearer {self.analyzer.model_api_key}",
                "Content-Type": "application/json"
            }
            
            data = json.dumps(payload).encode('utf-8')
            request = urllib.request.Request(url, data, headers)
            ssl_context = ssl._create_unverified_context()
            
            response = urllib.request.urlopen(request, context=ssl_context, timeout=30)
            response_data = json.loads(response.read().decode('utf-8'))
            
            if 'choices' in response_data and len(response_data['choices']) > 0:
                self.log(f"    ✓ {model_type} Model Test: PASSED (model responded)")
                return True
            else:
                self.log(f"    ✗ {model_type} Model Test: FAILED (unexpected response)")
                return False
                
        except urllib.error.HTTPError as e:
            error_body = e.read().decode('utf-8')
            self.log(f"    ✗ {model_type} Model Test: FAILED (HTTP {e.code})")
            self.log(f"      Error: {error_body[:150]}")
            return False
        except Exception as e:
            self.log(f"    ✗ {model_type} Model Test: FAILED ({str(e)[:100]})")
            return False
    
    def check_models_availability(self):
        """Check if configured models are available and working in Model Gateway"""
        if not self.analyzer.model_api_key:
            self.log("⚠ Cannot check models - no API key loaded")
            return False
        
        self.log("\n[Checking Model Availability...]")
        
        try:
            # Get Model Gateway URL from config
            gateway_url = self.config.get('model_gateway', {}).get('base_url',
                                                                    'https://model-gateway.gcldgenaigw.gc.micron.com/api/v1')
            url = f"{gateway_url}/models"
            
            headers = {
                "Authorization": f"Bearer {self.analyzer.model_api_key}",
                "Content-Type": "application/json"
            }
            
            request = urllib.request.Request(url, headers=headers)
            ssl_context = ssl._create_unverified_context()
            
            response = urllib.request.urlopen(request, context=ssl_context, timeout=30)
            data = json.loads(response.read().decode('utf-8'))
            
            if 'data' in data:
                available_models = [m.get('id') for m in data['data']]
                
                # Check configured models (from ai_modes.json)
                modes_data = self.load_modes_config()
                models = modes_data.get('models', {})
                analysis_model = models.get('analysis', {}).get('name', 'gemini-3.0-pro-preview')
                code_model = models.get('code_generation', {}).get('name', 'claude-sonnet-4-5')
                
                self.log(f"  Gateway: {gateway_url}")
                
                models_ok = True
                
                # Check Analysis Model
                if analysis_model in available_models:
                    self.log(f"  ✓ Analysis Model: {analysis_model} - LISTED")
                    # Test if it actually works
                    if not self.test_model_prompt(analysis_model, "Analysis"):
                        models_ok = False
                else:
                    self.log(f"  ✗ Analysis Model: {analysis_model} - NOT FOUND")
                    self.log(f"    Available: {', '.join(available_models[:5])}")
                    models_ok = False
                
                # Check Code Generation Model
                if code_model in available_models:
                    self.log(f"  ✓ Code Model: {code_model} - LISTED")
                    # Test if it actually works
                    if not self.test_model_prompt(code_model, "Code"):
                        models_ok = False
                else:
                    self.log(f"  ✗ Code Model: {code_model} - NOT FOUND")
                    self.log(f"    Available: {', '.join(available_models[:5])}")
                    models_ok = False
                
                return models_ok
            else:
                self.log(f"  ⚠ Unexpected response format from Model Gateway")
                return False
                
        except urllib.error.HTTPError as e:
            error_body = e.read().decode('utf-8')
            self.log(f"  ✗ HTTP Error {e.code}: {error_body[:200]}")
            return False
        except Exception as e:
            self.log(f"  ✗ Error checking models: {str(e)}")
            return False
    
    def check_all_connectivity(self):
        """Check all service connectivity"""
        self.log("\n" + "="*60)
        self.log("CONNECTIVITY CHECK")
        self.log("="*60)
        
        jira_ok = self.check_jira_connectivity()
        bb_ok = self.check_bitbucket_connectivity()
        models_ok = self.check_models_availability()
        
        self.log("\n" + "="*60)
        self.log("SUMMARY")
        self.log("="*60)
        self.log(f"  JIRA:      {'✓ OK' if jira_ok else '✗ FAILED'}")
        self.log(f"  Bitbucket: {'✓ OK' if bb_ok else '✗ FAILED'}")
        self.log(f"  Models:    {'✓ OK' if models_ok else '✗ FAILED'}")
        self.log("="*60)
        
        if not all([jira_ok, bb_ok, models_ok]):
            messagebox.showwarning("Connectivity Issues",
                "Some services are not accessible.\n\n"
                "Check the log for details and verify your credentials and configuration.")
        else:
            messagebox.showinfo("Connectivity Check",
                "All services are accessible!\n\n"
                "✓ JIRA\n✓ Bitbucket\n✓ Model Gateway")
    
    def load_credentials_with_lock(self):
        """Wrapper for load_credentials with GUI locking"""
        self.lock_gui()
        try:
            self.load_credentials()
        finally:
            self.unlock_gui()
    
    def check_all_connectivity_with_lock(self):
        """Wrapper for check_all_connectivity with GUI locking"""
        self.lock_gui()
        try:
            self.check_all_connectivity()
        finally:
            self.unlock_gui()
    
    def test_config_with_credential_check(self):
        """Test Config button handler - checks credentials are present before testing connectivity"""
        email = self.email_var.get().strip()
        jira_token = self.jira_token_var.get().strip()
        bb_token = self.bb_token_var.get().strip()
        model_key = self.model_key_var.get().strip()
        
        # Check if credentials are loaded/entered
        missing = []
        if not email:
            missing.append("Email")
        if not jira_token:
            missing.append("JIRA Token")
        if not bb_token:
            missing.append("Bitbucket Token")
        if not model_key:
            missing.append("Model API Key")
        
        if missing:
            messagebox.showerror(
                "Credentials Required",
                f"The following credentials are missing or empty:\n\n"
                f"  • " + "\n  • ".join(missing) + "\n\n"
                f"Please enter correct credentials (Load or enter manually) before testing configuration."
            )
            return
        
        # Auto-apply credentials before testing
        self.apply_credentials()
        
        # Now run connectivity check with lock
        self.check_all_connectivity_with_lock()
    
    def load_credentials(self):
        from tkinter import simpledialog
        
        # Load config first
        self.config = self.load_config()
        if self.config:
            self.jira_url_var.set(self.config.get('jira', {}).get('base_url', 'https://micron.atlassian.net'))
            self.bb_url_var.set(self.config.get('bitbucket', {}).get('base_url', 'https://bitbucket.micron.com/bbdc/scm'))
            self.project_key_var.set(self.config.get('bitbucket', {}).get('project_key', 'TESTSSD'))
            self.jira_project_var.set(self.config.get('jira', {}).get('project_key', 'TSESSD'))
            self.model_url_var.set(self.config.get('model_gateway', {}).get('base_url', 'https://model-gateway.gcldgenaigw.gc.micron.com/api/v1'))
            
            # Load model configuration from ai_modes.json
            modes_data = self.load_modes_config()
            models = modes_data.get('models', {})
            self.analysis_model_var.set(models.get('analysis', {}).get('name', 'gemini-3-pro-preview'))
            self.code_model_var.set(models.get('code_generation', {}).get('name', 'claude-sonnet-4-5'))
            
            self.log("✓ Configuration loaded from settings.json")
            self.log(f"  Analysis Model: {self.analysis_model_var.get()}")
            self.log(f"  Code Model: {self.code_model_var.get()}")
        
        # Load encrypted credentials
        password = simpledialog.askstring("Password", "Enter password:", show='*')
        if not password:
            return
        
        creds = CredentialManager.load_credentials(password)
        if creds:
            self.email_var.set(creds['email'])
            self.jira_token_var.set(creds['jira_token'])
            self.bb_token_var.set(creds['bitbucket_token'])
            self.model_key_var.set(creds['model_api_key'])
            # Also update URLs from credentials if they exist
            if 'jira_url' in creds:
                self.jira_url_var.set(creds['jira_url'])
            if 'bitbucket_url' in creds:
                self.bb_url_var.set(creds['bitbucket_url'])
            self.log("✓ Credentials loaded from encrypted file")
            self.apply_credentials()
            # Check all connectivity after credentials are loaded
            self.check_all_connectivity()
        else:
            messagebox.showerror("Error", "Failed to load credentials")
    
    def save_credentials(self):
        from tkinter import simpledialog
        
        # First save config
        self.save_config()
        
        # Then save encrypted credentials
        password = simpledialog.askstring("Password", "Enter password for encryption:", show='*')
        if not password:
            return
        
        confirm = simpledialog.askstring("Confirm", "Confirm password:", show='*')
        if password != confirm:
            messagebox.showerror("Error", "Passwords don't match")
            return
        
        email = self.email_var.get()
        if '@' in email:
            bb_user = email.split('@')[0]
        else:
            bb_user = email
        
        success = CredentialManager.save_credentials(
            email, bb_user,
            self.jira_token_var.get(),
            self.bb_token_var.get(),
            self.jira_url_var.get(),
            self.bb_url_var.get(),
            self.model_key_var.get(),
            password
        )
        
        if success:
            self.log("✓ Credentials saved to encrypted file")
            self.log("✓ Configuration saved to settings.json")
            # Auto-apply credentials after saving
            self.apply_credentials()
            messagebox.showinfo("Success", "Credentials and configuration saved and applied successfully")
        else:
            messagebox.showerror("Error", "Failed to save credentials")
    
    def apply_credentials_with_lock(self):
        """Wrapper for apply_credentials with GUI locking"""
        self.lock_gui()
        try:
            self.apply_credentials()
        finally:
            self.unlock_gui()
    
    def apply_credentials(self):
        email = self.email_var.get()
        if not email:
            messagebox.showerror("Error", "Email required")
            return
        
        # Apply credentials
        self.analyzer.email = email
        self.analyzer.bitbucket_username = email.split('@')[0] if '@' in email else email
        self.analyzer.jira_token = self.jira_token_var.get()
        self.analyzer.bitbucket_token = self.bb_token_var.get()
        self.analyzer.model_api_key = self.model_key_var.get()
        
        # Apply configuration from GUI
        self.analyzer.jira_base_url = self.jira_url_var.get()
        self.analyzer.bitbucket_base_url = self.bb_url_var.get()
        
        from jira_analyzer import AIGatewayClient
        self.analyzer.ai_client = AIGatewayClient(self.analyzer.model_api_key)
        
        # CRITICAL: Set debug mode IMMEDIATELY after creating AI client
        # This ensures debug output appears in GUI log during JIRA analysis
        debug_enabled = self.debug_var.get()
        if hasattr(self.analyzer.ai_client, 'config'):
            if 'settings' not in self.analyzer.ai_client.config:
                self.analyzer.ai_client.config['settings'] = {}
            self.analyzer.ai_client.config['settings']['debug'] = debug_enabled
        # Set the debug flag directly on the client
        self.analyzer.ai_client.debug = debug_enabled
        
        self.log("✓ Credentials and configuration applied")
        if debug_enabled:
            self.log("  ✓ DEBUG MODE ENABLED - Field extraction details will be shown")
        self.log(f"  JIRA: {self.analyzer.jira_base_url}")
        self.log(f"  Bitbucket: {self.analyzer.bitbucket_base_url}")
        self.log(f"  Project: {self.project_key_var.get()}")
        self.fetch_repos()
    
    def filter_repos(self, event=None):
        """Filter repository list based on typed text"""
        typed = self.repo_var.get().lower()
        if not typed:
            # Show all repos
            repo_names = [f"{r['slug']} - {r['name']}" for r in self.repos]
            self.repo_combo['values'] = repo_names
        else:
            # Filter repos
            filtered = [f"{r['slug']} - {r['name']}" for r in self.repos
                       if typed in r['slug'].lower() or typed in r['name'].lower()]
            self.repo_combo['values'] = filtered
    
    def filter_branches(self, event=None):
        """Filter branch list based on typed text"""
        typed = self.branch_var.get().lower()
        if not typed:
            # Show all branches
            self.branch_combo['values'] = self.branches
        else:
            # Filter branches
            filtered = [b for b in self.branches if typed in b.lower()]
            self.branch_combo['values'] = filtered
    
    def fetch_repos(self):
        if not self.analyzer.bitbucket_username:
            messagebox.showerror("Error", "Apply credentials first")
            return
        
        project_key = self.project_key_var.get()
        self.log(f"Fetching repositories from project: {project_key}...")
        repos = self.analyzer.list_project_repositories(project_key)
        
        if repos:
            self.repos = repos
            repo_names = [f"{r['slug']} - {r['name']}" for r in repos]
            self.repo_combo['values'] = repo_names
            self.log(f"✓ Found {len(repos)} repositories")
        else:
            self.log("✗ Failed to fetch repositories")
    
    def on_repo_selected(self, event):
        # Get selected repo slug from the combo text
        selected_text = self.repo_var.get()
        if ' - ' in selected_text:
            repo_slug = selected_text.split(' - ')[0]
        else:
            repo_slug = selected_text
        
        project_key = self.project_key_var.get()
        self.log(f"Fetching branches for {repo_slug}...")
        
        branches = self.analyzer.list_repository_branches(project_key, repo_slug)
        if branches:
            self.branches = branches
            self.branch_combo['values'] = branches
            self.log(f"✓ Found {len(branches)} branches")
            if branches:
                self.branch_combo.current(0)
        else:
            self.log("✗ Failed to fetch branches")
    
    def start_analysis(self):
        issue_key = self.issue_var.get().strip().upper()
        repo_text = self.repo_var.get().strip()
        branch = self.branch_var.get().strip()
        feature_branch_input = self.feature_branch_var.get().strip()
        
        if not all([issue_key, repo_text, branch]):
            messagebox.showerror("Error", "All fields required")
            return
        
        # Extract repo slug from selection (format: "slug - name")
        if ' - ' in repo_text:
            repo_slug = repo_text.split(' - ')[0].strip()
        else:
            # If user typed directly, use as-is
            repo_slug = repo_text
        
        self.log("\n" + "="*60)
        self.log("Starting JIRA Analysis...")
        self.log(f"Issue: {issue_key}")
        self.log(f"Repository: {repo_slug}")
        self.log(f"Branch: {branch}")
        self.log(f"Feature Branch Input: {feature_branch_input}")
        self.log(f"Debug - repo_text: {repo_text}")
        self.log(f"Debug - extracted repo_slug: {repo_slug}")
        self.log("="*60)
        
        # Determine feature branch name for confirmation message
        if feature_branch_input and "if empty," not in feature_branch_input.lower() and feature_branch_input.lower() != "auto populate":
            fb_display = feature_branch_input
        else:
            fb_display = f"feature/{issue_key}"
        
        # Confirm before starting
        response = messagebox.askyesno("Confirm",
            f"Start analysis for {issue_key}?\n\n"
            f"This will:\n"
            f"1. Clone {repo_slug}\n"
            f"2. Create/Checkout branch '{fb_display}'\n"
            f"3. Run AI analysis\n"
            f"4. Generate implementation plan\n\n"
            f"Continue?")
        
        if response:
            # Lock GUI during analysis
            self.lock_gui()
            
            # Run analysis in background thread
            thread = threading.Thread(target=self.run_analysis_with_unlock,
                                     args=(issue_key, repo_slug, branch, feature_branch_input))
            thread.daemon = True
            thread.start()
    
    def run_analysis_with_unlock(self, issue_key, repo_name, branch_name, feature_branch_input=None):
        """Wrapper for run_analysis that unlocks GUI when done"""
        try:
            self.run_analysis(issue_key, repo_name, branch_name, feature_branch_input)
        finally:
            # Unlock GUI when analysis completes or fails
            self.root.after(0, self.unlock_gui)
    
    def implement_code_changes_gui(self, repo_path, jira_analysis, impact_analysis, repo_index):
        """GUI version of implement_code_changes that uses messagebox instead of terminal input"""
        import os
        
        self.log("\n[Implementing Code Changes with AI]")
        
        # Get list of key files to analyze
        key_files = []
        for file_info in repo_index['file_index'][:20]:  # Top 20 files
            file_path = os.path.join(repo_path, file_info['path'])
            if os.path.isfile(file_path) and file_info['size'] < 100000:  # Skip large files
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                        key_files.append({
                            'path': file_info['path'],
                            'content': content[:2000]  # First 2000 chars
                        })
                except:
                    pass
        
        # Prepare implementation prompt
        impl_prompt = f"""Based on the JIRA analysis and code impact assessment, provide specific code changes:

**JIRA Analysis:**
{jira_analysis.get('analysis', 'N/A')}

**Impact Analysis:**
{impact_analysis.get('impact_analysis', 'N/A')}

**Sample Code Files:**
{json.dumps(key_files[:5], indent=2)}

Provide implementation in this EXACT format:

## FILES_TO_MODIFY
- path/to/file1.py: Brief description of changes
- path/to/file2.js: Brief description of changes

## CODE_CHANGES
### path/to/file1.py
```python
# Complete updated code for this file
```

### path/to/file2.js
```javascript
// Complete updated code for this file
```

## NEW_FILES
### path/to/newfile.py
```python
# Complete code for new file
```

## IMPLEMENTATION_GUIDE
Step-by-step guide for manual review and additional changes.
"""
        
        messages = [{"role": "user", "content": impl_prompt}]
        
        # Use code generation model
        result = self.analyzer.ai_client.chat_completion(messages, task_type="code_generation", mode="code_implementation")
        
        if result['success']:
            implementation_plan = result['response']['choices'][0]['message']['content']
            self.log("✓ AI generated implementation plan")
            
            # Save implementation plan
            plan_file = os.path.join(repo_path, "IMPLEMENTATION_PLAN.md")
            with open(plan_file, 'w', encoding='utf-8') as f:
                f.write(f"# Implementation Plan\n\n")
                f.write(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                f.write(implementation_plan)
            
            self.log(f"  Implementation plan saved to: IMPLEMENTATION_PLAN.md")
            
            # Ask via GUI messagebox instead of terminal
            apply_auto = messagebox.askyesno("Apply Code Changes",
                "Apply code changes automatically?\n\n"
                "This will modify files based on AI suggestions.\n"
                "You can review changes with 'git diff' afterwards.")
            
            if apply_auto:
                changes_applied = self.analyzer._apply_code_changes(repo_path, implementation_plan)
                if changes_applied > 0:
                    self.log(f"\n✓ Applied {changes_applied} code changes")
                    self.log("  Review changes with: git diff")
                else:
                    self.log("\n  No automatic changes applied. Review IMPLEMENTATION_PLAN.md for manual changes.")
            else:
                self.log("\n  Review IMPLEMENTATION_PLAN.md and apply changes manually.")
            
            return True
        else:
            self.log(f"✗ Failed to generate implementation plan: {result['error']}")
            return False
    
    def run_analysis(self, issue_key, repo_name, branch_name, feature_branch_input=None):
        """Run the full analysis in background thread"""
        import traceback
        try:
            self.log(f"\nDebug - run_analysis called with:")
            self.log(f"  issue_key: {issue_key}")
            self.log(f"  repo_name: {repo_name}")
            self.log(f"  branch_name: {branch_name}")
            self.log(f"  feature_branch_input: {feature_branch_input}")
            
            # Initialize workflow file
            self.init_workflow_file(issue_key)
            
            self.log("\n[Step 1] Fetching JIRA issue...")
            issue_data = self.analyzer.fetch_jira_issue(issue_key)
            if not issue_data:
                self.log("✗ Failed to fetch JIRA issue")
                messagebox.showerror("JIRA Error",
                    f"Failed to fetch JIRA issue {issue_key}.\n\nCheck the log for details.")
                return
            
            # Save JIRA issue data to workflow file
            fields = issue_data.get('fields', {})
            summary = fields.get('summary', 'N/A')
            description = fields.get('description', 'N/A')
            issue_text = f"Issue: {issue_key}\nSummary: {summary}\n\nDescription:\n{self.analyzer._extract_text_from_adf(description)}\n"
            self.save_workflow_step("ISSUE_KEY", issue_key)
            self.save_workflow_step("JIRA_ISSUE_DATA", issue_text)
            
            self.log("\n[Step 2] Analyzing JIRA request...")
            if self.debug_var.get():
                self.log("  [DEBUG] Using analysis model for JIRA analysis")
                self.log(f"  [DEBUG] Model: {self.analysis_model_var.get()}")
            jira_analysis = self.analyzer.analyze_jira_request(issue_data)
            if not jira_analysis.get('success'):
                error_msg = jira_analysis.get('error', 'Unknown error')
                self.log(f"✗ Failed to analyze JIRA request: {error_msg}")
                messagebox.showerror("AI Analysis Error",
                    f"Failed to analyze JIRA request.\n\nError: {error_msg}\n\n"
                    "Check your model configuration in ai_modes.json")
                return
            
            # Open chat for JIRA analysis review (don't save yet - wait for approval)
            self.log("\n[Interactive] Opening chat for JIRA analysis review...")
            self.root.after(0, lambda: self.create_interactive_chat(
                issue_key,
                "JIRA Analysis",
                jira_analysis.get('analysis', ''),
                lambda: self.finalize_step2_and_continue(issue_key, repo_name, branch_name, jira_analysis, feature_branch_input)
            ))
            
        except Exception as e:
            error_trace = traceback.format_exc()
            self.log(f"\n✗ Error: {str(e)}")
            self.log(f"\nFull Traceback:\n{error_trace}")
            messagebox.showerror("Error",
                f"Analysis failed: {str(e)}\n\nCheck log for full traceback")
    
    def finalize_step2_and_continue(self, issue_key, repo_name, branch_name, jira_analysis, feature_branch_input=None):
        """Finalize Step 2 (JIRA Analysis) after user approval, then continue"""
        # Save JIRA analysis to workflow file
        self.save_workflow_step("JIRA_ANALYSIS", jira_analysis.get('analysis', ''))
        
        # Log completion ONLY after user approval
        self.log("✓ JIRA Analysis complete and approved")
        
        # Continue to next step
        self.continue_analysis_step3(issue_key, repo_name, branch_name, jira_analysis, feature_branch_input)
    
    def continue_analysis_step3(self, issue_key, repo_name, branch_name, jira_analysis, feature_branch_input=None):
        """Continue with step 3 after user approves JIRA analysis"""
        import traceback
        try:
            
            self.log("\n[Step 3] Cloning repository...")
            repo_path = self.analyzer.clone_repository(repo_name, branch_name, issue_key)
            if not repo_path:
                self.log("✗ Failed to clone repository")
                messagebox.showerror("Repository Error",
                    f"Failed to clone repository {repo_name}.\n\nCheck the log for details.")
                return
            
            # Save repository info to workflow file
            repo_info = f"Repository: {repo_name}\nBase branch: {branch_name}\nLocal path: {repo_path}\n"
            self.save_workflow_step("REPOSITORY_PATH", repo_path)
            self.save_workflow_step("REPOSITORY_INFO", repo_info)
            
            # Auto-populate repository path to Implementation and Test Scenarios tabs
            self.impl_repo_var.set(repo_path)
            
            self.log("\n[Step 4] Creating feature branch...")
            if not self.analyzer.create_feature_branch(repo_path, issue_key, branch_name, feature_branch_input):
                self.log("✗ Failed to create feature branch")
                messagebox.showerror("Git Error",
                    f"Failed to create feature branch.\n\nCheck the log for details.")
                return
            
            # Update workflow file with feature branch info
            if feature_branch_input and feature_branch_input.strip() and "if empty," not in feature_branch_input.lower() and feature_branch_input.lower() != "auto populate":
                feature_branch_name = feature_branch_input.strip()
            else:
                feature_branch_name = f"feature/{issue_key}"
                
            self.feature_branch_var.set(feature_branch_name)
            repo_info_updated = f"Repository: {repo_name}\nBase branch: {branch_name}\nFeature branch: {feature_branch_name}\nLocal path: {repo_path}\n"
            self.save_workflow_step("REPOSITORY_INFO", repo_info_updated)
            
            # Ensure repository path is populated in Implementation and Test Scenarios tabs
            self.impl_repo_var.set(repo_path)
            
            self.log("\n[Step 5] Indexing repository...")
            repo_index = self.analyzer.index_repository(repo_path)
            
            self.log("\n[Step 6] Analyzing code impact...")
            if self.debug_var.get():
                self.log("  [DEBUG] Using analysis model for code impact")
                self.log(f"  [DEBUG] Model: {self.analysis_model_var.get()}")
            impact_analysis = self.analyzer.analyze_code_impact(repo_path, repo_index, jira_analysis)
            if not impact_analysis.get('success'):
                error_msg = impact_analysis.get('error', 'Unknown error')
                self.log(f"✗ Code impact analysis failed: {error_msg}")
                messagebox.showerror("AI Analysis Error",
                    f"Code impact analysis failed.\n\nError: {error_msg}")
                return
            
            # Open chat for impact analysis review (don't save yet - wait for approval)
            self.log("\n[Interactive] Opening chat for impact analysis review...")
            self.root.after(0, lambda: self.create_interactive_chat(
                issue_key,
                "Code Impact Analysis",
                impact_analysis.get('impact_analysis', ''),
                lambda: self.finalize_step6_and_continue(issue_key, repo_path, jira_analysis, impact_analysis, repo_index)
            ))
            
        except Exception as e:
            error_trace = traceback.format_exc()
            self.log(f"\n✗ Error: {str(e)}")
            self.log(f"\nFull Traceback:\n{error_trace}")
            messagebox.showerror("Error",
                f"Analysis failed: {str(e)}\n\nCheck log for full traceback")
    
    def finalize_step6_and_continue(self, issue_key, repo_path, jira_analysis, impact_analysis, repo_index):
        """Finalize Step 6 (Impact Analysis) after user approval, then continue"""
        # Save impact analysis to workflow file
        self.save_workflow_step("IMPACT_ANALYSIS", impact_analysis.get('impact_analysis', ''))
        
        # Log completion ONLY after user approval
        self.log("✓ Code Impact Analysis complete and approved")
        
        # Continue to next step
        self.continue_analysis_step7(issue_key, repo_path, jira_analysis, impact_analysis, repo_index)
    
    def continue_analysis_step7(self, issue_key, repo_path, jira_analysis, impact_analysis, repo_index):
        """Continue with remaining steps after user approves impact analysis"""
        import traceback
        try:
            
            self.log("\n[Step 7] Generating test scenarios...")
            if self.debug_var.get():
                self.log("  [DEBUG] Using code generation model for test scenarios")
                self.log(f"  [DEBUG] Model: {self.code_model_var.get()}")
            test_scenarios = self.analyzer.generate_test_scenarios(jira_analysis, repo_index)
            if not test_scenarios.get('success'):
                error_msg = test_scenarios.get('error', 'Unknown error')
                self.log(f"✗ Test scenario generation failed: {error_msg}")
                messagebox.showerror("AI Analysis Error",
                    f"Test scenario generation failed.\n\nError: {error_msg}")
                return
            
            # Save test scenarios to workflow file
            self.save_workflow_step("TEST_SCENARIOS", test_scenarios.get('test_scenarios', ''))
            
            self.log("\n[Step 8] Assessing risks...")
            if self.debug_var.get():
                self.log("  [DEBUG] Using analysis model for risk assessment")
                self.log(f"  [DEBUG] Model: {self.analysis_model_var.get()}")
            risk_assessment = self.analyzer.assess_risks(jira_analysis, impact_analysis)
            if not risk_assessment.get('success'):
                error_msg = risk_assessment.get('error', 'Unknown error')
                self.log(f"✗ Risk assessment failed: {error_msg}")
                messagebox.showerror("AI Analysis Error",
                    f"Risk assessment failed.\n\nError: {error_msg}")
                return
            
            # Save risk assessment to workflow file
            self.save_workflow_step("RISK_ASSESSMENT", risk_assessment.get('risk_assessment', ''))
            
            self.log("\n[Step 9] Generating implementation plan...")
            # Use GUI version that prompts via messagebox
            self.implement_code_changes_gui(repo_path, jira_analysis, impact_analysis, repo_index)
            
            self.log("\n[Step 10] Saving analysis report...")
            all_results = {
                'jira_analysis': jira_analysis,
                'impact_analysis': impact_analysis,
                'test_scenarios': test_scenarios,
                'risk_assessment': risk_assessment,
                'repo_path': repo_path,
                'repo_stats': repo_index['stats']
            }
            
            report_file = self.analyzer.save_analysis_report(issue_key, all_results)
            
            self.log("\n" + "="*60)
            self.log("✓ ANALYSIS COMPLETE!")
            self.log(f"✓ Repository: {repo_path}")
            self.log(f"✓ Report: {report_file}")
            self.log("="*60)
            
            # Open chat window for continued interaction
            self.root.after(100, lambda: self.create_chat_window(issue_key, jira_analysis))
            
            messagebox.showinfo("Success",
                f"Analysis complete!\n\n"
                f"Repository: {repo_path}\n"
                f"Report: {report_file}\n\n"
                f"A chat window has been opened for continued AI assistance.\n\n"
                f"Check the log for details.")
            
        except Exception as e:
            error_trace = traceback.format_exc()
            self.log(f"\n✗ Error: {str(e)}")
            self.log(f"\nFull Traceback:\n{error_trace}")
            messagebox.showerror("Error",
                f"Analysis failed: {str(e)}\n\nCheck log for full traceback")

    def create_interactive_chat(self, issue_key, step_name, analysis_result, continue_callback):
        """Create an interactive chat window for a specific analysis step"""
        # Check if chat window already exists and is valid
        try:
            if self.chat_window and self.chat_window.winfo_exists():
                self.chat_window.destroy()
        except:
            pass
        
        self.chat_window = tk.Toplevel(self.root)
        self.chat_window.title(f"AI Chat - {issue_key} - {step_name}")
        self.chat_window.geometry("900x700")
        
        # Make window modal
        self.chat_window.transient(self.root)
        self.chat_window.grab_set()
        self._centre_dialog(self.chat_window, 900, 700)
        
        self.log(f"✓ Opening interactive chat for {step_name}...")
        
        # Initialize chat with context from this step
        self.current_chat_messages = [
            {"role": "system", "content": f"You are helping with JIRA issue {issue_key}, specifically reviewing the {step_name}."},
            {"role": "assistant", "content": f"**{step_name} Result:**\n\n{analysis_result}"}
        ]
        
        # Header with step info
        header_frame = ttk.Frame(self.chat_window, padding="10")
        header_frame.pack(fill=tk.X)
        
        ttk.Label(header_frame, text=f"Step: {step_name}", font=('Arial', 12, 'bold')).pack(anchor=tk.W)
        
        # Step-specific instructions
        if "JIRA Analysis" in step_name or "Analyze JIRA" in step_name:
            instruction_text = "Review the analysis and chat with AI. If possible, feed the AI with Micron specific data such as density, step for more accurate output.\nClick 'Approve & Continue' when satisfied."
        elif "Implementation" in step_name or "Code" in step_name:
            instruction_text = "Review the code implementation and chat with AI. If incorrect, tell the AI which file specifically need to change and continuously feedback for better code.\nClick 'Approve & Continue' when satisfied."
        elif "Test Scenarios" in step_name:
            instruction_text = "Review the test scenarios and chat with AI. Request AI to add specific test cases and to remove those deemed unnecessary.\nClick 'Approve & Continue' when satisfied."
        elif "Risk" in step_name or "Validation" in step_name:
            instruction_text = "Review the risk assessment and validation plan. Ask AI to elaborate on specific risks or suggest additional mitigation strategies.\nClick 'Approve & Continue' when satisfied."
        else:
            instruction_text = "Review the analysis and chat with AI for refinements.\nClick 'Approve & Continue' when satisfied."
        
        ttk.Label(header_frame, text=instruction_text,
                 foreground='blue').pack(anchor=tk.W, pady=5)
        
        # Chat display area
        chat_frame = ttk.LabelFrame(self.chat_window, text="AI Conversation", padding="10")
        chat_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.chat_display = scrolledtext.ScrolledText(chat_frame, height=25, width=90, wrap=tk.WORD)
        self.chat_display.pack(fill=tk.BOTH, expand=True)
        self.chat_display.config(state=tk.DISABLED)
        
        # Input area
        input_frame = ttk.Frame(self.chat_window, padding="10")
        input_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        ttk.Label(input_frame, text="Your message:").pack(anchor=tk.W)
        
        self.chat_input = scrolledtext.ScrolledText(input_frame, height=3, width=90, wrap=tk.WORD)
        self.chat_input.pack(fill=tk.X, pady=5)
        
        # Bind Enter key (Shift+Enter for new line)
        self.chat_input.bind('<Return>', lambda e: self.send_chat_message() if not e.state & 0x1 else None)
        
        btn_frame = ttk.Frame(input_frame)
        btn_frame.pack(fill=tk.X)
        
        ttk.Button(btn_frame, text="Send Message", command=self.send_chat_message).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="Clear Chat", command=self.clear_chat).pack(side=tk.LEFT, padx=2)
        
        # Approve button to continue to next step
        approve_btn = ttk.Button(btn_frame, text="✓ Approve & Continue", 
                                command=lambda: self.approve_and_continue(continue_callback))
        approve_btn.pack(side=tk.RIGHT, padx=2)
        
        ttk.Button(btn_frame, text="Cancel Analysis", 
                  command=lambda: self.cancel_analysis()).pack(side=tk.RIGHT, padx=2)
        
        # Add initial analysis result
        self.add_chat_message("assistant", f"**{step_name} Complete!**\n\n{analysis_result}\n\n---\n\nYou can:\n• Ask questions about this analysis\n• Request clarifications or modifications\n• Ask me to refine specific sections\n\nWhen satisfied, click 'Approve & Continue' to save the final version.")
        
        self.log(f"\n✓ Interactive chat opened for {step_name}")
        self.log(f"  User can refine analysis before approval")
    
    def approve_and_continue(self, continue_callback):
        """User approved the current step, continue to next"""
        if self.chat_window:
            self.chat_window.destroy()
            self.chat_window = None
        
        self.log("✓ User approved - continuing to next step...")
        
        # Run the continuation callback in a thread
        thread = threading.Thread(target=continue_callback)
        thread.daemon = True
        thread.start()
    
    def cancel_analysis(self):
        """User cancelled the analysis"""
        if messagebox.askyesno("Cancel Analysis", "Are you sure you want to cancel the analysis?"):
            if self.chat_window:
                self.chat_window.destroy()
                self.chat_window = None
            self.log("\n✗ Analysis cancelled by user")


    def create_chat_window(self, issue_key, jira_analysis):
        """Create a chat window for continued AI interaction after analysis"""
        # Check if chat window already exists and is valid
        try:
            if self.chat_window and self.chat_window.winfo_exists():
                self.chat_window.lift()
                self.log(f"✓ Chat window already open for {issue_key}")
                return
        except:
            pass
        
        self.chat_window = tk.Toplevel(self.root)
        self.chat_window.title(f"AI Chat - {issue_key}")
        self.chat_window.geometry("800x600")
        self.chat_window.transient(self.root)
        self._centre_dialog(self.chat_window, 800, 600)
        
        self.log(f"✓ Creating chat window for {issue_key}...")
        
        # Initialize chat with context from analysis
        self.current_chat_messages = [
            {"role": "system", "content": f"You are helping with JIRA issue {issue_key}. Context: {jira_analysis.get('analysis', '')[:500]}"},
        ]
        
        # Chat display area
        chat_frame = ttk.LabelFrame(self.chat_window, text="AI Conversation", padding="10")
        chat_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        self.chat_display = scrolledtext.ScrolledText(chat_frame, height=25, width=90, wrap=tk.WORD)
        self.chat_display.pack(fill=tk.BOTH, expand=True)
        self.chat_display.config(state=tk.DISABLED)
        
        # Input area
        input_frame = ttk.Frame(self.chat_window, padding="10")
        input_frame.pack(fill=tk.X, padx=10, pady=(0, 10))
        
        ttk.Label(input_frame, text="Your message:").pack(anchor=tk.W)
        
        self.chat_input = scrolledtext.ScrolledText(input_frame, height=3, width=90, wrap=tk.WORD)
        self.chat_input.pack(fill=tk.X, pady=5)
        
        # Bind Enter key (Shift+Enter for new line)
        self.chat_input.bind('<Return>', lambda e: self.send_chat_message() if not e.state & 0x1 else None)
        
        btn_frame = ttk.Frame(input_frame)
        btn_frame.pack(fill=tk.X)
        
        ttk.Button(btn_frame, text="Send Message", command=self.send_chat_message).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="Clear Chat", command=self.clear_chat).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="Close", command=self.chat_window.destroy).pack(side=tk.RIGHT, padx=2)
        
        # Add initial message
        self.add_chat_message("assistant", f"Hello! I've analyzed JIRA issue {issue_key}. How can I help you further with this implementation?")
        
        self.log(f"\n✓ Chat window opened for {issue_key}")
    
    def add_chat_message(self, role, content):
        """Add a message to the chat display"""
        self.chat_display.config(state=tk.NORMAL)
        
        if role == "user":
            self.chat_display.insert(tk.END, "\n[You]\n", "user_tag")
            self.chat_display.tag_config("user_tag", foreground="blue", font=("Arial", 10, "bold"))
        else:
            self.chat_display.insert(tk.END, "\n[AI Assistant]\n", "ai_tag")
            self.chat_display.tag_config("ai_tag", foreground="green", font=("Arial", 10, "bold"))
        
        self.chat_display.insert(tk.END, f"{content}\n")
        self.chat_display.see(tk.END)
        self.chat_display.config(state=tk.DISABLED)
    
    def send_chat_message(self):
        """Send a message and get AI response with conversation history"""
        message = self.chat_input.get("1.0", tk.END).strip()
        if not message:
            return
        
        # Add user message to display and history
        self.add_chat_message("user", message)
        self.current_chat_messages.append({"role": "user", "content": message})
        
        # Clear input
        self.chat_input.delete("1.0", tk.END)
        
        # Show typing indicator
        self.chat_display.config(state=tk.NORMAL)
        self.chat_display.insert(tk.END, "\n[AI is typing...]\n", "typing_tag")
        self.chat_display.tag_config("typing_tag", foreground="gray", font=("Arial", 9, "italic"))
        self.chat_display.see(tk.END)
        self.chat_display.config(state=tk.DISABLED)
        self.chat_window.update()
        
        # Get AI response with full conversation history
        try:
            # Determine which model to use based on the chat window title
            chat_title = self.chat_window.title() if self.chat_window else ""
            
            # Use appropriate task type based on the step
            # Implementation/Code steps use code_generation model (which now handles both code gen and implementation impact)
            if "Implementation" in chat_title or "Code" in chat_title:
                task_type = "code_generation"
            else:
                task_type = "analysis"
            
            if self.debug_var.get():
                model_config = self.analyzer.ai_client.get_model_config(task_type)
                model_name = model_config.get('name', 'unknown')
                self.log(f"  [DEBUG] Chat using {task_type} model: {model_name}")
                self.log(f"  [DEBUG] Message count in history: {len(self.current_chat_messages)}")
            
            result = self.analyzer.ai_client.chat_completion(
                self.current_chat_messages,  # Send entire conversation history
                task_type=task_type,
                mode="chat_continuation"
            )
            
            # Remove typing indicator
            self.chat_display.config(state=tk.NORMAL)
            self.chat_display.delete("end-2l", "end-1l")
            self.chat_display.config(state=tk.DISABLED)
            
            if result['success']:
                response_text = result['response']['choices'][0]['message']['content']
                self.add_chat_message("assistant", response_text)
                
                # Add AI response to history for context continuation
                self.current_chat_messages.append({"role": "assistant", "content": response_text})
                
                self.log(f"✓ Chat message processed (history: {len(self.current_chat_messages)} messages)")
            else:
                error_msg = f"Error: {result.get('error', 'Unknown error')}"
                self.add_chat_message("assistant", error_msg)
                self.log(f"✗ Chat error: {result.get('error')}")
        
        except Exception as e:
            # Remove typing indicator
            self.chat_display.config(state=tk.NORMAL)
            self.chat_display.delete("end-2l", "end-1l")
            self.chat_display.config(state=tk.DISABLED)
            
            error_msg = f"Error communicating with AI: {str(e)}"
            self.add_chat_message("assistant", error_msg)
            self.log(f"✗ Chat exception: {str(e)}")
    
    def clear_chat(self):
        """Clear chat history but keep initial context"""
        if messagebox.askyesno("Clear Chat", "Clear conversation history?\n\nThis will reset the chat but keep the JIRA context."):
            # Keep only the system message
            system_msg = self.current_chat_messages[0] if self.current_chat_messages else None
            self.current_chat_messages = [system_msg] if system_msg else []
            
            self.chat_display.config(state=tk.NORMAL)
            self.chat_display.delete("1.0", tk.END)
            self.chat_display.config(state=tk.DISABLED)
            
            self.add_chat_message("assistant", "Chat cleared. How can I help you?")
            self.log("✓ Chat history cleared")


def main():
    root = tk.Tk()
    app = SimpleGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()