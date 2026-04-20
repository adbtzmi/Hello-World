# -*- coding: utf-8 -*-
"""
Validation Document Populator
Automatically populates validation document template with checkout test results.
"""

import os
import sys
import re
from typing import Dict, List, Optional, Tuple
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Handle imports for both standalone and module usage
try:
    from model.analyzers.trace_analyzer import TraceAnalysis, TraceAnalyzer
except ModuleNotFoundError:
    # Running as standalone script, add parent directory to path
    sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..'))
    from model.analyzers.trace_analyzer import TraceAnalysis, TraceAnalyzer


class ValidationDocumentPopulator:
    """
    Populates validation document template with checkout test results.
    
    Uses trace analysis data to fill in:
    - Test execution summary
    - Pass/fail statistics
    - Test time information
    - Trace file references
    - Database references
    """
    
    def __init__(self, template_path: str = "template_validation.docx"):
        """
        Initialize the populator.
        
        Args:
            template_path: Path to the validation document template
        """
        self.template_path = template_path
    
    def populate_document(
        self,
        trace_analysis: TraceAnalysis,
        output_path: str,
        jira_key: Optional[str] = None,
        mid: Optional[str] = None,
        sku: Optional[str] = None,
        trace_file_path: Optional[str] = None,
        database_path: Optional[str] = None,
        scenario: Optional[str] = None,
        old_flow_time: Optional[float] = None,
    ) -> bool:
        """
        Populate validation document with trace analysis results.
        
        Args:
            trace_analysis: TraceAnalysis object with test results
            output_path: Path to save the populated document
            jira_key: JIRA issue key (e.g., "TSESSD-14270")
            mid: Module ID
            sku: SKU identifier
            trace_file_path: Path to trace file for reference
            database_path: Path to database file for reference
            scenario: Test scenario description
            old_flow_time: Old flow test time for comparison (hours)
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Load template
            doc = Document(self.template_path)
            
            # Replace placeholders in document
            self._replace_text_in_document(doc, trace_analysis, jira_key, mid, sku)
            
            # Populate test validation section
            self._populate_test_validation(
                doc,
                trace_analysis,
                trace_file_path,
                database_path,
                scenario,
                old_flow_time,
                is_passing=True
            )
            
            # If there are failed tests, populate failing case section
            if trace_analysis.failed_tests > 0:
                self._populate_test_validation(
                    doc,
                    trace_analysis,
                    trace_file_path,
                    database_path,
                    scenario,
                    old_flow_time,
                    is_passing=False
                )
            
            # Save populated document
            doc.save(output_path)
            return True
            
        except Exception as e:
            import logging
            logging.getLogger("bento_app").error(
                f"Error populating validation document: {e}")
            print(f"Error populating validation document: {e}")
            return False
    
    def _replace_text_in_document(
        self,
        doc: Document,
        trace_analysis: TraceAnalysis,
        jira_key: Optional[str],
        mid: Optional[str],
        sku: Optional[str]
    ):
        """Replace placeholder text throughout the document."""
        replacements = {
            'TSESSD-XXXXX': jira_key or 'TSESSD-XXXXX',
            'ABCD': jira_key or 'ABCD',
        }
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    self._replace_text_in_paragraph(paragraph, old_text, new_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            if old_text in paragraph.text:
                                self._replace_text_in_paragraph(paragraph, old_text, new_text)
    
    def _replace_text_in_paragraph(self, paragraph, old_text: str, new_text: str):
        """Replace text in a paragraph while preserving formatting."""
        if old_text in paragraph.text:
            # Simple replacement - preserves basic formatting
            inline = paragraph.runs
            for run in inline:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
    
    def _populate_test_validation(
        self,
        doc: Document,
        trace_analysis: TraceAnalysis,
        trace_file_path: Optional[str],
        database_path: Optional[str],
        scenario: Optional[str],
        old_flow_time: Optional[float],
        is_passing: bool = True
    ):
        """
        Populate test validation section (PASSING or FAILING case).
        
        Args:
            doc: Document object
            trace_analysis: TraceAnalysis object
            trace_file_path: Path to trace file
            database_path: Path to database file
            scenario: Test scenario description
            old_flow_time: Old flow test time (hours)
            is_passing: True for PASSING CASE, False for FAILING CASE
        """
        section_name = "PASSING CASE" if is_passing else "FAILING CASE"
        
        # Find the section in the document
        section_found = False
        for i, paragraph in enumerate(doc.paragraphs):
            if section_name in paragraph.text:
                section_found = True
                # Look for the table after this section
                # In Word documents, tables are separate from paragraphs
                break
        
        if not section_found:
            return
        
        # Find tables and populate them
        for table in doc.tables:
            # Check if this table contains test validation data
            # Look for specific markers like "SKU:", "MID:", "Test Time"
            table_text = '\n'.join([cell.text for row in table.rows for cell in row.cells])
            
            if 'SKU:' in table_text and 'MID:' in table_text:
                self._populate_summary_table(table, trace_analysis, is_passing)
            
            if 'Test Time (hr)' in table_text:
                self._populate_test_time_table(table, trace_analysis, old_flow_time)
            
            if 'Trace' in table_text and 'Database' in table_text:
                self._populate_file_references_table(
                    table, trace_file_path, database_path
                )
    
    def _populate_summary_table(
        self,
        table,
        trace_analysis: TraceAnalysis,
        is_passing: bool
    ):
        """Populate SKU/MID/SUMMARY table."""
        # Generate summary text
        if is_passing:
            summary = (
                f"[PASS] All tests passed ({trace_analysis.passed_tests}/{trace_analysis.total_tests})\n"
                f"Pass Rate: {trace_analysis.pass_rate:.2f}%\n"
                f"Duration: {trace_analysis.duration_seconds / 3600:.2f} hours"
            )
        else:
            summary = (
                f"[FAIL] {trace_analysis.failed_tests} test(s) failed\n"
                f"Pass Rate: {trace_analysis.pass_rate:.2f}%\n"
                f"Failed Tests:\n"
            )
            # Add first few failed tests
            failed_tests = [t for t in trace_analysis.test_results if t.result == 'FAIL']
            for test in failed_tests[:5]:  # Show first 5 failed tests
                summary += f"  - {test.test_name}\n"
            if len(failed_tests) > 5:
                summary += f"  ... and {len(failed_tests) - 5} more\n"
        
        # Find and populate SUMMARY cell
        for row in table.rows:
            for cell in row.cells:
                if 'SUMMARY:' in cell.text:
                    # Get the next cell (value cell)
                    cell_idx = row.cells.index(cell)
                    if cell_idx + 1 < len(row.cells):
                        value_cell = row.cells[cell_idx + 1]
                        value_cell.text = summary
    
    def _populate_test_time_table(
        self,
        table,
        trace_analysis: TraceAnalysis,
        old_flow_time: Optional[float]
    ):
        """Populate test time comparison table."""
        new_flow_time = trace_analysis.duration_seconds / 3600  # Convert to hours
        
        # Calculate delta if old flow time is provided
        if old_flow_time:
            delta_hours = new_flow_time - old_flow_time
            delta_percent = (delta_hours / old_flow_time) * 100 if old_flow_time > 0 else 0
        else:
            old_flow_time = 0
            delta_hours = 0
            delta_percent = 0
        
        # Find and populate cells
        for row in table.rows:
            row_text = ' '.join([cell.text for cell in row.cells])
            
            if 'Test Time (hr)' in row_text:
                # This is the header row, next row should have values
                continue
            
            # Look for data row (usually has numbers)
            cells = row.cells
            if len(cells) >= 4:
                # Assuming structure: New flow | Old flow | TT Delta | TT Delta %
                # Try to find the right cells
                for i, cell in enumerate(cells):
                    cell_text = cell.text.strip()
                    
                    # Check if this looks like a time value cell
                    if cell_text and (
                        re.match(r'^\d+\.?\d*$', cell_text) or
                        cell_text in ['New flow', 'Old flow', 'TT Delta']
                    ):
                        # Found a data row
                        if i + 3 < len(cells):
                            cells[i].text = f"{new_flow_time:.5f}"
                            cells[i + 1].text = f"{old_flow_time:.5f}"
                            cells[i + 2].text = f"{delta_hours:.7f}"
                            cells[i + 3].text = f"{delta_percent:.2f}%"
                            break
    
    def _populate_file_references_table(
        self,
        table,
        trace_file_path: Optional[str],
        database_path: Optional[str]
    ):
        """Populate trace and database file reference table."""
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                
                if cell_text == 'Trace' and trace_file_path:
                    # Add trace file path to next cell or same cell
                    if i + 1 < len(row.cells):
                        row.cells[i + 1].text = trace_file_path
                    else:
                        cell.text = f"Trace\n{trace_file_path}"
                
                elif cell_text == 'Database' and database_path:
                    # Add database path to next cell or same cell
                    if i + 1 < len(row.cells):
                        row.cells[i + 1].text = database_path
                    else:
                        cell.text = f"Database\n{database_path}"
    
    def populate_from_trace_file(
        self,
        trace_file_path: str,
        output_path: str,
        **kwargs
    ) -> bool:
        """
        Convenience method to analyze trace file and populate document in one step.
        
        Args:
            trace_file_path: Path to trace file
            output_path: Path to save populated document
            **kwargs: Additional parameters (jira_key, mid, sku, etc.)
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Analyze trace file
            analyzer = TraceAnalyzer()
            trace_analysis = analyzer.analyze_file(trace_file_path)
            
            # Populate document
            return self.populate_document(
                trace_analysis=trace_analysis,
                output_path=output_path,
                trace_file_path=trace_file_path,
                **kwargs
            )
        except Exception as e:
            print(f"Error in populate_from_trace_file: {e}")
            return False
    
    def batch_populate(
        self,
        trace_files: List[str],
        output_dir: str,
        **common_kwargs
    ) -> Dict[str, bool]:
        """
        Populate multiple validation documents from multiple trace files.
        
        Args:
            trace_files: List of trace file paths
            output_dir: Directory to save populated documents
            **common_kwargs: Common parameters for all documents
            
        Returns:
            Dictionary mapping trace file path to success status
        """
        results = {}
        
        for trace_file in trace_files:
            try:
                # Generate output filename
                trace_basename = os.path.basename(trace_file)
                output_filename = trace_basename.replace('.txt', '_validation.docx')
                output_path = os.path.join(output_dir, output_filename)
                
                # Populate document
                success = self.populate_from_trace_file(
                    trace_file_path=trace_file,
                    output_path=output_path,
                    **common_kwargs
                )
                
                results[trace_file] = success
                
            except Exception as e:
                print(f"Error processing {trace_file}: {e}")
                results[trace_file] = False
        
        return results


# Convenience function
def populate_validation_doc(
    trace_file_path: str,
    output_path: str,
    template_path: str = "template_validation.docx",
    **kwargs
) -> bool:
    """
    Quick function to populate validation document from trace file.
    
    Args:
        trace_file_path: Path to trace file
        output_path: Path to save populated document
        template_path: Path to template document
        **kwargs: Additional parameters (jira_key, mid, sku, etc.)
        
    Returns:
        True if successful, False otherwise
    """
    populator = ValidationDocumentPopulator(template_path)
    return populator.populate_from_trace_file(trace_file_path, output_path, **kwargs)


if __name__ == '__main__':
    # Example usage
    import sys
    
    if len(sys.argv) < 3:
        print("Usage: python validation_doc_populator.py <trace_file> <output_doc> [jira_key] [mid] [sku]")
        sys.exit(1)
    
    trace_file = sys.argv[1]
    output_doc = sys.argv[2]
    jira_key = sys.argv[3] if len(sys.argv) > 3 else None
    mid = sys.argv[4] if len(sys.argv) > 4 else None
    sku = sys.argv[5] if len(sys.argv) > 5 else None
    
    print(f"Populating validation document...")
    print(f"  Trace file: {trace_file}")
    print(f"  Output: {output_doc}")
    if jira_key:
        print(f"  JIRA: {jira_key}")
    if mid:
        print(f"  MID: {mid}")
    if sku:
        print(f"  SKU: {sku}")
    print()
    
    success = populate_validation_doc(
        trace_file_path=trace_file,
        output_path=output_doc,
        jira_key=jira_key,
        mid=mid,
        sku=sku
    )
    
    if success:
        print(f"[SUCCESS] Validation document created successfully: {output_doc}")
    else:
        print(f"[FAILED] Failed to create validation document")
        sys.exit(1)
